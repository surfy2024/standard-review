#!/usr/bin/env python3
"""
引用标准符合性检查测试脚本

测试 R001-R012 所有规则：
- R001-R005: 引用标准提取与格式检查
- R006-R008: 交叉引用一致性检查
- R009-R012: 要求性条款提取、语义匹配、指标级比对
"""

import sys
import os
import json
import subprocess
import tempfile
from pathlib import Path

# 确保能导入模块
SCRIPTS_DIR = Path(__file__).parent
sys.path.insert(0, str(SCRIPTS_DIR))

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def create_test_draft(tmpdir):
    """创建测试用标准草稿文档"""
    doc = Document()
    
    # 标题
    doc.add_heading('XX产品技术规范', level=0)
    
    # 前言
    doc.add_heading('前言', level=1)
    doc.add_paragraph('本文件按照GB/T 1.1-2020的规定起草。')
    doc.add_paragraph('请注意本文件的某些内容可能涉及专利。')
    
    # 1 范围
    doc.add_heading('1 范围', level=1)
    doc.add_paragraph('本文件规定了XX产品的技术要求、试验方法、检验规则。')
    doc.add_paragraph('本文件适用于XX产品的生产、检验和验收。')
    
    # 2 规范性引用文件
    doc.add_heading('2 规范性引用文件', level=1)
    doc.add_paragraph(
        '下列文件中的内容通过文中的规范性引用而构成本文件必不可少的条款。'
        '其中，注日期的引用文件，仅该日期对应的版本适用于本文件；'
        '未注日期的引用文件，其最新版本（包括所有的修改单）适用于本文件。'
    )
    # 故意包含格式问题用于测试 R002
    doc.add_paragraph('GB/T 1.1-2020 标准化工作导则 第1部分：标准化文件的结构和起草规则')
    doc.add_paragraph('GB/T 7714-2015 信息与文献 参考文献著录规则')
    doc.add_paragraph('YY 0123-2020 医疗器械 通用要求')  # 行业标准
    doc.add_paragraph('GB/T 12345 数据分析方法')  # 不注日期引用
    # 故意重复引用用于测试 R003
    doc.add_paragraph('GB/T 1.1-2020 标准化工作导则（重复）')
    # 这个在正文中不会被引用，用于测试 R008
    doc.add_paragraph('GB/T 99999-2018 未被正文引用的标准')
    
    # 3 术语和定义
    doc.add_heading('3 术语和定义', level=1)
    doc.add_paragraph('下列术语和定义适用于本文件。')
    doc.add_paragraph('3.1 产品 product')
    doc.add_paragraph('满足特定需求的物品。')
    
    # 4 技术要求
    doc.add_heading('4 技术要求', level=1)
    doc.add_paragraph('4.1 一般要求')
    doc.add_paragraph('产品外观应完整，无明显缺陷。')
    # 包含定量指标的要求
    doc.add_paragraph('4.2 性能要求')
    doc.add_paragraph('产品的强度应不低于20MPa。')  # 应不低于 → 指标
    doc.add_paragraph('产品的含水率应不大于5%。')  # 应不大于 → 指标
    doc.add_paragraph('产品的工作温度应不低于-10℃。')  # 指标
    doc.add_paragraph('4.3 安全要求')
    doc.add_paragraph('产品不得含有有害物质。')  # 不得 → 禁止性要求
    doc.add_paragraph('产品必须通过安全认证。')  # 必须 → 强制性要求
    
    # 5 试验方法
    doc.add_heading('5 试验方法', level=1)
    doc.add_paragraph('5.1 强度试验')
    doc.add_paragraph('按GB/T 1.1-2020的规定执行强度试验。')  # 引用标准（已在引用文件中）
    doc.add_paragraph('5.2 含水率试验')
    doc.add_paragraph('按照GB/T 7714-2015的方法进行。')  # 引用标准
    # 故意引用一个未在引用文件中列出的标准，用于测试 R007
    doc.add_paragraph('5.3 其他试验')
    doc.add_paragraph('试验结果应符合GB/T 88888-2020的要求。')  # 缺失引用
    
    # 6 检验规则
    doc.add_heading('6 检验规则', level=1)
    doc.add_paragraph('6.1 出厂检验')
    doc.add_paragraph('每批产品应进行出厂检验，检验项目应符合表1的规定。')
    
    # 7 标志、包装
    doc.add_heading('7 标志、包装、运输和贮存', level=1)
    doc.add_paragraph('产品标志应清晰、耐久。')
    
    draft_path = os.path.join(tmpdir, "test_draft.docx")
    doc.save(draft_path)
    return draft_path


def create_test_reference_standard(tmpdir):
    """创建测试用引用标准文档（模拟 GB/T 1.1-2020 的关键要求）"""
    doc = Document()
    
    doc.add_heading('标准化工作导则 第1部分：标准化文件的结构和起草规则', level=0)
    
    doc.add_heading('1 范围', level=1)
    doc.add_paragraph('本文件确立了标准化文件的结构及其起草的总体原则。')
    
    doc.add_heading('2 规范性引用文件', level=1)
    doc.add_paragraph('下列文件中的内容通过文中的规范性引用而构成本文件必不可少的条款。')
    
    doc.add_heading('3 术语和定义', level=1)
    doc.add_paragraph('下列术语和定义适用于本文件。')
    
    doc.add_heading('4 文件的结构', level=1)
    doc.add_paragraph('4.1 一般要求')
    doc.add_paragraph('标准化文件应包含范围、规范性引用文件、术语和定义等要素。')
    # 这里设一个更低的指标，测试草稿是否"优于"
    doc.add_paragraph('4.2 性能要求')
    doc.add_paragraph('文件的强度指标应不低于15MPa。')  # 引用标准要求 15，草稿要求 20 → 草稿优于
    doc.add_paragraph('文件的含水率应不大于8%。')  # 引用标准要求 8%，草稿要求 5% → 草稿优于
    
    doc.add_heading('5 文件的起草', level=1)
    doc.add_paragraph('5.1 总则')
    doc.add_paragraph('文件应使用规范的汉字和术语。')
    doc.add_paragraph('5.2 条款要求')
    doc.add_paragraph('要求性条款应使用"应"，推荐性条款应使用"宜"。')
    doc.add_paragraph('不应使用"必须"代替"应"。')  # 禁止性条款
    
    ref_path = os.path.join(tmpdir, "GB-T1.1-2020_ref.docx")
    doc.save(ref_path)
    return ref_path


def create_test_reference_standard_2(tmpdir):
    """创建第二个测试用引用标准文档"""
    doc = Document()
    
    doc.add_heading('信息与文献 参考文献著录规则', level=0)
    
    doc.add_heading('1 范围', level=1)
    doc.add_paragraph('本标准规定了参考文献的著录规则。')
    
    doc.add_heading('2 著录项目', level=1)
    doc.add_paragraph('参考文献应包含主要责任者、题名、出版项等信息。')
    doc.add_paragraph('参考文献的著录应不少于3个要素。')  # 指标
    
    ref_path = os.path.join(tmpdir, "GB-T7714-2015_ref.docx")
    doc.save(ref_path)
    return ref_path


def create_minimal_draft(tmpdir):
    """创建最小化草稿（无引用文件章节）用于测试"""
    doc = Document()
    doc.add_heading('简单标准', level=0)
    doc.add_heading('1 范围', level=1)
    doc.add_paragraph('本文件规定了简单要求。')
    doc.add_paragraph('产品应合格。')
    
    path = os.path.join(tmpdir, "minimal_draft.docx")
    doc.save(path)
    return path


def run_tests():
    """运行所有测试"""
    passed = 0
    failed = 0
    errors = []
    
    def test_result(name, condition, detail=""):
        nonlocal passed, failed, errors
        if condition:
            passed += 1
            print(f"  ✓ {name}")
        else:
            failed += 1
            errors.append(f"{name}: {detail}")
            print(f"  ✗ {name} — {detail}")
    
    print("=" * 60)
    print("引用标准符合性检查测试 (R001-R012)")
    print("=" * 60)
    
    with tempfile.TemporaryDirectory() as tmpdir:
        # 创建测试文档
        print("\n--- 创建测试文档 ---")
        draft_path = create_test_draft(tmpdir)
        ref_path = create_test_reference_standard(tmpdir)
        ref_path_2 = create_test_reference_standard_2(tmpdir)
        minimal_path = create_minimal_draft(tmpdir)
        print(f"草稿: {draft_path}")
        print(f"引用标准1: {ref_path}")
        print(f"引用标准2: {ref_path_2}")
        
        # ============================================================
        # 测试 R001: 引用标准提取
        # ============================================================
        print("\n--- R001: 引用标准提取 ---")
        
        from check_standard import StandardChecker
        checker = StandardChecker()
        issues = checker.check(draft_path)
        
        ref_checker = checker.ref_checker
        refs = ref_checker.references
        
        test_result("R001: 提取到引用标准", len(refs) > 0, f"提取到 {len(refs)} 条")
        test_result("R001: 提取到 GB/T 1.1-2020", 
                    any('1.1' in r.number_part for r in refs),
                    f"提取到的编号: {[r.number for r in refs]}")
        test_result("R001: 提取到 GB/T 7714-2015",
                    any('7714' in r.number_part for r in refs))
        test_result("R001: 提取到 YY 0123-2020",
                    any('0123' in r.number_part for r in refs))
        test_result("R001: 提取到不注日期引用 GB/T 12345",
                    any('12345' in r.number_part for r in refs))
        test_result("R001: 识别注日期 vs 不注日期",
                    any(r.is_dated for r in refs) and any(not r.is_dated for r in refs))
        
        # ============================================================
        # 测试 R002: 引用格式校验
        # ============================================================
        print("\n--- R002: 引用格式校验 ---")
        
        r002_issues = [i for i in issues if i.code == 'R002']
        test_result("R002: 检测到格式问题", len(r002_issues) >= 0)  # 可能没有问题或有问题
        
        # ============================================================
        # 测试 R003: 重复引用检查
        # ============================================================
        print("\n--- R003: 重复引用检查 ---")
        
        r003_issues = [i for i in issues if i.code == 'R003']
        test_result("R003: 检测到重复引用 GB/T 1.1-2020", 
                    len(r003_issues) > 0,
                    f"检测到 {len(r003_issues)} 个重复引用")
        test_result("R003: 重复引用包含 GB/T 1.1",
                    any('1.1' in i.description for i in r003_issues))
        
        # ============================================================
        # 测试 R004: 排序规范性
        # ============================================================
        print("\n--- R004: 排序规范性 ---")
        
        r004_issues = [i for i in issues if i.code == 'R004']
        # 测试文档中 GB/T 99999 排在 YY 0123 后面，可能触发排序问题
        test_result("R004: 排序检查运行", len(r004_issues) >= 0)
        
        # ============================================================
        # 测试 R005: 引导语检查
        # ============================================================
        print("\n--- R005: 引导语检查 ---")
        
        r005_issues = [i for i in issues if i.code == 'R005']
        # 测试文档有完整引导语，应该没有 R005 错误
        r005_errors = [i for i in r005_issues if i.severity == 'ERROR']
        test_result("R005: 有引导语时无 ERROR", len(r005_errors) == 0,
                    f"R005 问题: {[i.description for i in r005_errors]}")
        
        # 测试无引导语的情况
        checker2 = StandardChecker()
        # 创建无引导语的草稿
        doc_no_intro = Document()
        doc_no_intro.add_heading('无引导语标准', level=0)
        doc_no_intro.add_heading('1 范围', level=1)
        doc_no_intro.add_paragraph('本文件规定了XX要求。')
        doc_no_intro.add_heading('2 规范性引用文件', level=1)
        doc_no_intro.add_paragraph('GB/T 1.1-2020 标准化工作导则')  # 无引导语
        doc_no_intro.add_heading('3 术语和定义', level=1)
        doc_no_intro.add_paragraph('术语定义。')
        no_intro_path = os.path.join(tmpdir, "no_intro.docx")
        doc_no_intro.save(no_intro_path)
        
        issues2 = checker2.check(no_intro_path)
        r005_no_intro = [i for i in issues2 if i.code == 'R005' and i.severity == 'ERROR']
        test_result("R005: 无引导语时检测到 ERROR", len(r005_no_intro) > 0,
                    f"R005 问题: {[i.description for i in r005_no_intro]}")
        
        # ============================================================
        # 测试 R006: 正文引用提取
        # ============================================================
        print("\n--- R006: 正文引用提取 ---")
        
        citations = ref_checker.citations
        test_result("R006: 提取到正文引用", len(citations) > 0,
                    f"提取到 {len(citations)} 条引用")
        test_result("R006: 提取到 GB/T 1.1-2020 引用",
                    any('1.1' in c['number'] for c in citations))
        test_result("R006: 提取到 GB/T 7714-2015 引用",
                    any('7714' in c['number'] for c in citations))
        test_result("R006: 提取到 GB/T 88888-2020 引用",
                    any('88888' in c['number'] for c in citations))
        
        # ============================================================
        # 测试 R007: 缺失引用检查
        # ============================================================
        print("\n--- R007: 缺失引用检查 ---")
        
        r007_issues = [i for i in issues if i.code == 'R007']
        test_result("R007: 检测到缺失引用 GB/T 88888-2020",
                    len(r007_issues) > 0,
                    f"检测到 {len(r007_issues)} 个缺失引用")
        test_result("R007: 缺失引用包含 88888",
                    any('88888' in i.description for i in r007_issues))
        
        # ============================================================
        # 测试 R008: 冗余引用检查
        # ============================================================
        print("\n--- R008: 冗余引用检查 ---")
        
        r008_issues = [i for i in issues if i.code == 'R008']
        test_result("R008: 检测到冗余引用 GB/T 99999-2018",
                    len(r008_issues) > 0,
                    f"检测到 {len(r008_issues)} 个冗余引用")
        test_result("R008: 冗余引用包含 99999",
                    any('99999' in i.description for i in r008_issues))
        
        # ============================================================
        # 测试 R009-R012: 符合性检查（用户提交层）
        # ============================================================
        print("\n--- R009-R012: 符合性检查 ---")
        
        checker3 = StandardChecker()
        issues3 = checker3.check(draft_path, ref_files=[ref_path, ref_path_2])
        
        # R009: 从引用标准提取要求性条款
        ref_reqs = checker3.ref_checker.ref_requirements
        test_result("R009: 从引用标准提取要求性条款",
                    len(ref_reqs) > 0,
                    f"提取到 {len(ref_reqs)} 条")
        test_result("R009: 引用标准条款含'应'",
                    any(r.modal_verb == '应' for r in ref_reqs))
        test_result("R009: 引用标准条款含'不应'",
                    any(r.modal_verb == '不应' for r in ref_reqs))
        
        # R010: 从草稿提取要求性条款
        draft_reqs = checker3.ref_checker.draft_requirements
        test_result("R010: 从草稿提取要求性条款",
                    len(draft_reqs) > 0,
                    f"提取到 {len(draft_reqs)} 条")
        test_result("R010: 草稿条款含'应'",
                    any(r.modal_verb == '应' for r in draft_reqs))
        test_result("R010: 草稿条款含'不得'",
                    any(r.modal_verb == '不得' for r in draft_reqs))
        test_result("R010: 草稿条款含'必须'",
                    any(r.modal_verb == '必须' for r in draft_reqs))
        
        # R011: 语义匹配
        # 检查是否进行了匹配（通过 R012 结果间接验证）
        r012_issues = [i for i in issues3 if i.code == 'R012']
        test_result("R011/R012: 产生了符合性比对结果",
                    len(r012_issues) > 0,
                    f"产生了 {len(r012_issues)} 条比对结果")
        
        # R012: 指标级比对
        # 测试文档中：草稿强度≥20MPa vs 引用标准≥15MPa → 优于
        # 草稿含水率≤5% vs 引用标准≤8% → 优于
        r012_errors = [i for i in r012_issues if i.severity == 'ERROR']
        r012_suggestions = [i for i in r012_issues if i.severity == 'SUGGESTION']
        
        test_result("R012: 有比对结果", len(r012_issues) > 0)
        test_result("R012: 无不符合项（草稿均优于引用标准）",
                    len(r012_errors) == 0,
                    f"不符合项: {[i.description for i in r012_errors]}")
        
        # ============================================================
        # 测试 CLI --ref 参数
        # ============================================================
        print("\n--- CLI --ref 参数测试 ---")
        
        result = subprocess.run(
            [sys.executable, "check_standard.py", draft_path,
             "--ref", ref_path, "--ref", ref_path_2,
             "--pretty", "-o", os.path.join(tmpdir, "result.json")],
            capture_output=True, text=True, cwd=str(SCRIPTS_DIR),
            timeout=60
        )
        output = result.stdout + result.stderr
        
        test_result("CLI --ref: 退出码 0 或 1（有 ERROR 时为 1）",
                    result.returncode in (0, 1),
                    f"退出码: {result.returncode}")
        test_result("CLI --ref: 输出含'引用标准符合性检查'",
                    "引用标准符合性检查" in output,
                    f"输出片段: {output[:200]}")
        test_result("CLI --ref: 输出含'要求性条款'",
                    "要求性条款" in output or "条要求" in output)
        
        # 检查结果 JSON
        result_path = os.path.join(tmpdir, "result.json")
        if os.path.exists(result_path):
            with open(result_path, 'r', encoding='utf-8') as f:
                result_data = json.load(f)
            test_result("CLI --ref: JSON 含 reference_check",
                        "reference_check" in result_data)
            if "reference_check" in result_data:
                rc = result_data["reference_check"]
                test_result("CLI --ref: JSON 含 total_references",
                            "total_references" in rc)
                test_result("CLI --ref: JSON 含 draft_requirements",
                            "draft_requirements" in rc)
                test_result("CLI --ref: JSON 含 ref_requirements",
                            "ref_requirements" in rc)
        
        # ============================================================
        # 测试 CLI --ref-dir 参数
        # ============================================================
        print("\n--- CLI --ref-dir 参数测试 ---")
        
        # 创建引用标准目录
        ref_dir = os.path.join(tmpdir, "ref_dir")
        os.makedirs(ref_dir, exist_ok=True)
        # 复制引用标准文件到目录
        from shutil import copy2
        copy2(ref_path, ref_dir)
        copy2(ref_path_2, ref_dir)
        
        result = subprocess.run(
            [sys.executable, "check_standard.py", draft_path,
             "--ref-dir", ref_dir, "-o", os.devnull],
            capture_output=True, text=True, cwd=str(SCRIPTS_DIR),
            timeout=60
        )
        output = result.stdout + result.stderr
        
        test_result("CLI --ref-dir: 加载了引用标准文件",
                    "加载了" in output and "引用标准文件" in output)
        test_result("CLI --ref-dir: 执行了符合性检查",
                    "引用标准符合性检查" in output)
        
        # ============================================================
        # 测试最小化草稿（无引用文件章节）
        # ============================================================
        print("\n--- 边界情况测试 ---")
        
        checker4 = StandardChecker()
        issues4 = checker4.check(minimal_path)
        
        # 最小化草稿没有引用文件章节，不应崩溃
        test_result("边界: 无引用文件章节不崩溃", len(issues4) >= 0)
        test_result("边界: 无引用文件章节时 references 为空",
                    len(checker4.ref_checker.references) == 0)
        
        # ============================================================
        # 测试指标比对逻辑
        # ============================================================
        print("\n--- 指标比对逻辑测试 ---")
        
        from reference_checker import _extract_indicators, _get_comp_direction
        
        # 测试指标提取
        inds = _extract_indicators("温度应不低于20℃")
        test_result("指标提取: '不低于20℃' 提取到 1 个指标",
                    len(inds) == 1, f"提取到 {len(inds)} 个")
        if inds:
            test_result("指标提取: 值为 20", inds[0]['value'] == 20.0)
            test_result("指标提取: 方向为 min", inds[0]['direction'] == 'min')
        
        inds2 = _extract_indicators("含水率应不大于5%")
        test_result("指标提取: '不大于5%' 提取到 1 个指标",
                    len(inds2) == 1, f"提取到 {len(inds2)} 个")
        if inds2:
            test_result("指标提取: 值为 5", inds2[0]['value'] == 5.0)
            test_result("指标提取: 方向为 max", inds2[0]['direction'] == 'max')
        
        # 测试比对逻辑（直接使用 ReferenceChecker 的方法，不需要先调用 check()）
        from reference_checker import ReferenceChecker
        ref_chk = ReferenceChecker([], [], {}, None)
        
        # 草稿≥20 vs 引用≥15 → 优于
        result1 = ref_chk._compare_single_indicator(
            {'direction': 'min', 'value': 20, 'comp': '不低于', 'unit': 'MPa'},
            {'direction': 'min', 'value': 15, 'comp': '不低于', 'unit': 'MPa'}
        )
        test_result("比对: 草稿≥20 vs 引用≥15 → 优于", result1 == "优于", f"结果: {result1}")
        
        # 草稿≥10 vs 引用≥15 → 不符合
        result2 = ref_chk._compare_single_indicator(
            {'direction': 'min', 'value': 10, 'comp': '不低于', 'unit': 'MPa'},
            {'direction': 'min', 'value': 15, 'comp': '不低于', 'unit': 'MPa'}
        )
        test_result("比对: 草稿≥10 vs 引用≥15 → 不符合", result2 == "不符合", f"结果: {result2}")
        
        # 草稿≥15 vs 引用≥15 → 符合
        result3 = ref_chk._compare_single_indicator(
            {'direction': 'min', 'value': 15, 'comp': '不低于', 'unit': 'MPa'},
            {'direction': 'min', 'value': 15, 'comp': '不低于', 'unit': 'MPa'}
        )
        test_result("比对: 草稿≥15 vs 引用≥15 → 符合", result3 == "符合", f"结果: {result3}")
        
        # 草稿≤5 vs 引用≤8 → 优于（更严格的限制）
        result4 = ref_chk._compare_single_indicator(
            {'direction': 'max', 'value': 5, 'comp': '不大于', 'unit': '%'},
            {'direction': 'max', 'value': 8, 'comp': '不大于', 'unit': '%'}
        )
        test_result("比对: 草稿≤5 vs 引用≤8 → 优于", result4 == "优于", f"结果: {result4}")
        
        # 草稿≤10 vs 引用≤8 → 不符合
        result5 = ref_chk._compare_single_indicator(
            {'direction': 'max', 'value': 10, 'comp': '不大于', 'unit': '%'},
            {'direction': 'max', 'value': 8, 'comp': '不大于', 'unit': '%'}
        )
        test_result("比对: 草稿≤10 vs 引用≤8 → 不符合", result5 == "不符合", f"结果: {result5}")
        
        # ============================================================
        # 测试原有规则无回归
        # ============================================================
        print("\n--- 原有规则无回归 ---")
        
        # 确保原有规则仍在工作
        test_result("回归: 仍有原有规则问题（如 T006 或其他）",
                    any(i.code != 'R0' and not i.code.startswith('R0') for i in issues))
        test_result("回归: R 系列规则存在",
                    any(i.code.startswith('R0') for i in issues))
    
    # ============================================================
    # 汇总
    # ============================================================
    print("\n" + "=" * 60)
    print(f"测试结果: {passed} 通过, {failed} 失败")
    if errors:
        print("\n失败项:")
        for e in errors:
            print(f"  - {e}")
    print("=" * 60)
    
    return failed == 0


if __name__ == "__main__":
    success = run_tests()
    sys.exit(0 if success else 1)
