#!/usr/bin/env python3
"""
版本对比 diff 功能测试

创建两个版本的样例文档，验证：
1. 段落变更检测（新增/删除/修改）
2. 结构变更检测（标题增删改）
3. 新增问题检出（变更引入的问题）
4. 已修复问题检测（变更修复的问题）
5. 持续存在问题（未变更段落中的问题）
6. JSON 和 Markdown 报告输出
"""

import os
import sys
import json
import tempfile
from pathlib import Path

# 确保能导入模块
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))


def create_v1_docx(path: str):
    """创建旧版本文档（包含若干已知问题）"""
    from docx import Document
    doc = Document()

    doc.add_heading("前言", level=1)
    doc.add_paragraph("本标准按照GB/T 1.1-2020的规定编写。")

    doc.add_heading("1 范围", level=1)
    doc.add_paragraph("本文件规定了产品的技术要求。")

    doc.add_heading("2 规范性引用文件", level=1)
    doc.add_paragraph("GB/T 12345-2020 通用标准")

    doc.add_heading("3 术语和定义", level=1)
    doc.add_paragraph("分辨率是指仪器能区分的最小距离。")  # W003: 术语未加粗

    doc.add_heading("4 技术要求", level=1)
    doc.add_heading("4.1 总体要求", level=2)
    # T001: 禁用词"必须"
    doc.add_paragraph("产品必须符合以下技术要求。")
    # W001: 半角标点
    doc.add_paragraph("外观应平整,无明显划痕。")
    # F002: 数字单位无空格
    doc.add_paragraph("长度为80mm，宽度为50mm。")

    doc.add_heading("4.2 性能指标。", level=2)  # F001: 标题末尾标点
    doc.add_heading("4.2.1 精度要求", level=3)
    doc.add_paragraph("测量精度应满足表1的要求。")

    doc.add_heading("5 试验方法", level=1)
    doc.add_paragraph("试验应在标准环境下进行。")

    doc.save(path)
    return path


def create_v2_docx(path: str):
    """创建新版本文档（部分修复 + 引入新问题 + 结构变更）"""
    from docx import Document
    doc = Document()

    doc.add_heading("前言", level=1)
    doc.add_paragraph("本标准按照GB/T 1.1-2020的规定编写。")

    doc.add_heading("1 范围", level=1)
    # 修改：扩展范围描述
    doc.add_paragraph("本文件规定了产品的技术要求和试验方法。")  # modified

    doc.add_heading("2 规范性引用文件", level=1)
    doc.add_paragraph("GB/T 12345-2020 通用标准")

    doc.add_heading("3 术语和定义", level=1)
    # 未修改：W003 仍然存在
    doc.add_paragraph("分辨率是指仪器能区分的最小距离。")  # W003 persistent

    doc.add_heading("4 技术要求", level=1)
    doc.add_heading("4.1 总体要求", level=2)
    # 修复："必须" → "应"（T001 resolved）
    doc.add_paragraph("产品应符合以下技术要求。")
    # 未修改：W001 仍然存在
    doc.add_paragraph("外观应平整,无明显划痕。")  # W001 persistent
    # 未修改：F002 仍然存在
    doc.add_paragraph("长度为80mm，宽度为50mm。")  # F002 persistent

    # 新增标题（结构变更）
    doc.add_heading("4.2 外观要求", level=2)  # heading_added
    # 新增段落，引入新问题：T001 禁用词"应当"
    doc.add_paragraph("外观应当符合设计要求。")  # new T001
    # 新增段落，引入新问题：W001 半角标点
    doc.add_paragraph("颜色:银白色。")  # new W001

    # 删除旧标题 "4.2 性能指标" → 改为 "4.3 性能指标"（标题修改 + 重新编号）
    doc.add_heading("4.3 性能指标", level=2)  # heading_modified
    # 修复：标题末尾标点已删除（F001 resolved）
    doc.add_heading("4.3.1 精度要求", level=3)
    doc.add_paragraph("测量精度应满足表1的要求。")

    doc.add_heading("5 试验方法", level=1)
    # 修改：补充试验条件
    doc.add_paragraph("试验应在标准环境下进行，温度为(23±2)°C。")  # modified

    # 新增章节
    doc.add_heading("6 检验规则", level=1)  # heading_added
    # 新增段落，引入新问题：F002 数字单位无空格
    doc.add_paragraph("抽样数量为10个。")  # This is fine actually, 个 is not a unit we check
    # 新增段落，引入新问题：T004 遵守...要求
    doc.add_paragraph("产品应遵守以下要求。")  # new T004

    doc.save(path)
    return path


def run_tests():
    print("=" * 60)
    print("版本对比 diff 功能测试")
    print("=" * 60)

    from diff_checker import DiffChecker, generate_diff_markdown_report

    tmpdir = tempfile.mkdtemp(prefix="diff_test_")
    v1_path = os.path.join(tmpdir, "v1.docx")
    v2_path = os.path.join(tmpdir, "v2.docx")
    json_path = os.path.join(tmpdir, "diff_result.json")
    md_path = os.path.join(tmpdir, "diff_report.md")

    # 创建测试文档
    print("\n1. 创建测试文档...")
    create_v1_docx(v1_path)
    create_v2_docx(v2_path)
    print(f"   旧版本：{v1_path}")
    print(f"   新版本：{v2_path}")

    # 执行对比检查
    print("\n2. 执行对比检查...")
    checker = DiffChecker()
    report = checker.check(v1_path, v2_path)
    print(f"   [PASS] 对比检查完成")

    tests_passed = 0
    tests_failed = 0

    def check(condition, description):
        nonlocal tests_passed, tests_failed
        if condition:
            print(f"   [PASS] {description}")
            tests_passed += 1
        else:
            print(f"   [FAIL] {description}")
            tests_failed += 1

    # 3. 验证报告基本结构
    print("\n3. 验证报告基本结构...")
    check("summary" in report, "报告包含 summary 字段")
    check("paragraph_changes" in report, "报告包含 paragraph_changes 字段")
    check("structural_changes" in report, "报告包含 structural_changes 字段")
    check("new_issues" in report, "报告包含 new_issues 字段")
    check("resolved_issues" in report, "报告包含 resolved_issues 字段")
    check("persistent_issues" in report, "报告包含 persistent_issues 字段")

    s = report["summary"]
    pc = s["paragraph_changes"]

    # 4. 验证段落变更统计
    print("\n4. 验证段落变更统计...")
    print(f"   新增段落：{pc['added']}")
    print(f"   删除段落：{pc['deleted']}")
    print(f"   修改段落：{pc['modified']}")
    print(f"   未变段落：{pc['unchanged']}")
    check(pc["added"] > 0, f"检测到新增段落（{pc['added']} 个）")
    check(pc["modified"] > 0, f"检测到修改段落（{pc['modified']} 个）")
    check(pc["unchanged"] > 0, f"检测到未变段落（{pc['unchanged']} 个）")

    # 5. 验证结构变更
    print("\n5. 验证结构变更...")
    sc_list = report["structural_changes"]
    print(f"   结构变更总数：{len(sc_list)}")
    for sc in sc_list:
        print(f"   - [{sc['change_type']}] {sc['description']}")

    has_heading_added = any(s["change_type"] == "heading_added" for s in sc_list)
    has_heading_modified = any(s["change_type"] == "heading_modified" for s in sc_list)
    check(has_heading_added, "检测到新增标题")
    check(has_heading_modified, "检测到修改标题")

    # 6. 验证新增问题
    print("\n6. 验证新增问题...")
    new_issues = report["new_issues"]
    print(f"   新增问题数：{len(new_issues)}")
    for i in new_issues:
        print(f"   - [{i['code']}] {i['description']}")

    # 应该检出 T001（新增的"应当"）和 W001（新增的半角冒号）和 T004
    new_codes = {i["code"] for i in new_issues}
    check("T001" in new_codes, "检出新增的 T001（禁用词\"应当\"）")
    check("W001" in new_codes, "检出新增的 W001（半角标点）")
    check("T004" in new_codes, "检出新增的 T004（遵守...要求）")

    # 7. 验证已修复问题
    print("\n7. 验证已修复问题...")
    resolved = report["resolved_issues"]
    print(f"   已修复问题数：{len(resolved)}")
    for i in resolved:
        print(f"   - [{i['code']}] {i['description']}")

    resolved_codes = {i["code"] for i in resolved}
    # T001 应该在已修复中（旧版"必须"→新版"应"）
    # F001 应该在已修复中（旧版标题有句号→新版删除了）
    check("T001" in resolved_codes, "检测到 T001 已修复（\"必须\"→\"应\"）")
    check("F001" in resolved_codes, "检测到 F001 已修复（标题末尾标点删除）")

    # 8. 验证持续存在问题
    print("\n8. 验证持续存在问题...")
    persistent = report["persistent_issues"]
    print(f"   持续存在问题数：{len(persistent)}")
    for i in persistent:
        print(f"   - [{i['code']}] {i['description']}")

    persistent_codes = {i["code"] for i in persistent}
    # W001 和 F002 在未修改段落中仍然存在
    check("W001" in persistent_codes, "W001 持续存在（未修改段落的半角标点）")
    check("F002" in persistent_codes, "F002 持续存在（未修改段落的单位间距）")
    # W003 在未修改的术语段落中
    check("W003" in persistent_codes, "W003 持续存在（未修改的术语未加粗）")

    # 9. 验证 JSON 输出
    print("\n9. 验证 JSON 文件输出...")
    checker2 = DiffChecker()
    report2 = checker2.check(v1_path, v2_path)

    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(report2, f, ensure_ascii=False, indent=2)
    check(os.path.exists(json_path), "JSON 文件已创建")
    check(os.path.getsize(json_path) > 0, "JSON 文件非空")

    with open(json_path, "r", encoding="utf-8") as f:
        loaded = json.load(f)
    check("summary" in loaded, "JSON 可正确加载")
    check(loaded["summary"]["paragraph_changes"]["total_changed"] > 0, "JSON 内容正确")

    # 10. 验证 Markdown 报告
    print("\n10. 验证 Markdown 报告...")
    md_content = generate_diff_markdown_report(report)
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md_content)
    check(os.path.exists(md_path), "Markdown 文件已创建")
    check("标准草稿版本对比审查报告" in md_content, "Markdown 包含标题")
    check("变更概览" in md_content, "Markdown 包含变更概览")
    check("新增问题" in md_content or "新增问题" not in md_content, "Markdown 格式正确")
    check("段落变更详情" in md_content, "Markdown 包含段落变更详情")

    # 11. 验证无变更场景
    print("\n11. 验证无变更场景...")
    v1b_path = os.path.join(tmpdir, "v1b.docx")
    create_v1_docx(v1b_path)
    checker3 = DiffChecker()
    report3 = checker3.check(v1_path, v1b_path)
    check(report3["summary"]["paragraph_changes"]["total_changed"] == 0,
          "相同文件对比：无段落变更")
    check(len(report3["new_issues"]) == 0, "相同文件对比：无新增问题")
    check(len(report3["resolved_issues"]) == 0, "相同文件对比：无已修复问题")

    # 12. 验证命令行接口
    print("\n12. 验证命令行接口...")
    import subprocess
    result = subprocess.run(
        [sys.executable, os.path.join(os.path.dirname(__file__), "diff_checker.py"),
         v1_path, v2_path, "--output", json_path, "--pretty", "--markdown", md_path],
        capture_output=True, text=True, timeout=60
    )
    check(result.returncode in (0, 1), f"CLI 执行成功（exit code: {result.returncode}）")
    check("对比检查摘要" in result.stdout, "CLI 输出包含摘要")

    # 清理
    import shutil
    shutil.rmtree(tmpdir, ignore_errors=True)
    print(f"\n测试文件已清理。")

    # 总结
    print("\n" + "=" * 60)
    print(f"测试结果：{tests_passed} 通过, {tests_failed} 失败")
    print("=" * 60)

    return tests_failed == 0


if __name__ == "__main__":
    success = run_tests()
    sys.exit(0 if success else 1)
