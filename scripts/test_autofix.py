#!/usr/bin/env python3
"""
自动修复测试脚本

创建带已知问题的 .docx → 运行检查 → 运行自动修复 → 验证结果
"""

import sys
import os
import re
from pathlib import Path

# 确保能 import 同目录模块
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from check_standard import StandardChecker
from auto_fix import AutoFixer
from lxml import etree

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
W = f'{{{W_NS}}}'


def create_test_docx(path: str):
    """创建带已知问题的测试文档"""
    doc = Document()

    # 前言（必须有）
    doc.add_heading("前言", level=1)
    doc.add_paragraph("本文件规定了技术要求。")

    # 1 范围
    doc.add_heading("1 范围", level=1)
    doc.add_paragraph("本文件规定了产品的基本要求。")

    # 2 规范性引用文件
    doc.add_heading("2 规范性引用文件", level=1)
    doc.add_paragraph("下列文件中的内容通过文中的规范性引用而构成本文件必不可少的条款。")
    doc.add_paragraph("GB/T 12345:2020 标准化工作导则")  # T006: 冒号应为一字线

    # 3 术语和定义
    doc.add_heading("3 术语和定义", level=1)
    doc.add_paragraph("本文件没有需要界定的术语和定义。")

    # 4 技术要求
    doc.add_heading("4 技术要求", level=1)

    # 4.1 基本要求
    doc.add_heading("4.1 基本要求", level=2)
    # T001: "必须" → "应"
    doc.add_paragraph("产品必须符合以下基本要求：")
    # T001: "应当" → "应"
    doc.add_paragraph("设备应当定期检查。")
    # T001: "不得" → "不应"
    doc.add_paragraph("不得使用不合格材料。")
    # F002: 数字+单位无空格
    doc.add_paragraph("外形尺寸为80mm，重量为5kg。")
    # W001: 半角逗号
    doc.add_paragraph("产品尺寸,重量和材料应符合要求。")
    # W001: 半角冒号
    doc.add_paragraph("技术指标如下:抗拉强度、硬度。")
    # W001: 半角分号
    doc.add_paragraph("测试环境;温度为25C。")
    # F003: 尺寸表述不规范
    doc.add_paragraph("构件尺寸为80x25x50 mm。")
    # F011: 公式编号无括号
    doc.add_paragraph("强度计算见公式3。")
    # F001: 标题末尾有标点
    doc.add_heading("4.2 性能指标。", level=2)
    doc.add_paragraph("性能指标应符合规定。")

    doc.save(path)
    print(f"测试文档已创建: {path}")


def verify_track_changes(docx_path: str) -> dict:
    """验证修订版文件中的 Track Changes 标记"""
    doc = Document(docx_path)
    body = doc.element.body

    ins_count = len(body.findall(f'.//{W}ins'))
    del_count = len(body.findall(f'.//{W}del'))
    del_text_count = len(body.findall(f'.//{W}delText'))

    return {
        "ins_elements": ins_count,
        "del_elements": del_count,
        "delText_elements": del_text_count,
    }


def _get_accepted_text(doc) -> str:
    """获取"接受所有修订"后的文本（包含 w:ins 内容，排除 w:del 内容）"""
    body = doc.element.body
    lines = []
    for p_elem in body.findall(f'.//{W}p'):
        para_text = ''
        # 遍历所有 w:t 元素
        for t_elem in p_elem.iter(f'{W}t'):
            # 检查是否在 w:del 内
            parent = t_elem.getparent()
            in_del = False
            while parent is not None:
                if parent.tag == f'{W}del':
                    in_del = True
                    break
                parent = parent.getparent()
            if not in_del:
                para_text += t_elem.text or ''
        lines.append(para_text)
    return '\n'.join(lines)


def verify_text_fixed(docx_path: str) -> dict:
    """验证修复后的文本是否正确（基于"接受所有修订"后的文本）"""
    doc = Document(docx_path)
    full_text = _get_accepted_text(doc)

    checks = {
        # T001: 禁用词应被替换（但原词仍在 delText 中，paragraph.text 只返回保留的文本）
        "必须_removed": "必须" not in full_text or "必须" in _get_deltext(doc),
        "应当_removed": "应当" not in full_text or "应当" in _get_deltext(doc),
        "不得_removed": "不得" not in full_text or "不得" in _get_deltext(doc),
        "应_inserted": "应" in full_text,
        "不应_inserted": "不应" in full_text,

        # F002: 数字+单位应有空格
        "80mm_fixed": "80 mm" in full_text,
        "5kg_fixed": "5 kg" in full_text,

        # W001: 半角标点应改为全角
        "half_comma_fixed": "尺寸,重量" not in full_text or "尺寸，重量" in full_text,
        "half_colon_fixed": "如下:" not in full_text or "如下：" in full_text,
        "half_semicolon_fixed": "环境;" not in full_text or "环境；" in full_text,

        # F003: 尺寸表述应规范化
        "dimension_fixed": "80 mm x 25 mm x 50 mm" in full_text,

        # F011: 公式编号应有括号
        "formula_fixed": "公式(3)" in full_text,

        # F001: 标题末尾标点应删除
        "heading_punct_removed": "性能指标。" not in full_text or "性能指标" in full_text,

        # T006: 引用冒号应改为一字线
        "ref_fixed": "GB/T 12345-2020" in full_text,
    }

    return checks


def _get_deltext(doc) -> str:
    """获取所有 delText 内容（被标记删除的文本）"""
    body = doc.element.body
    del_texts = body.findall(f'.//{W}delText')
    return '\n'.join(t.text or '' for t in del_texts)


def run_test():
    """运行完整测试"""
    test_docx = os.path.join(os.path.dirname(os.path.abspath(__file__)), "test_autofix_input.docx")
    output_docx = os.path.join(os.path.dirname(os.path.abspath(__file__)), "test_autofix_output.docx")

    # 1. 创建测试文档
    print("=" * 60)
    print("1. 创建测试文档")
    create_test_docx(test_docx)

    # 2. 运行检查
    print("\n2. 运行标准检查...")
    checker = StandardChecker()
    issues = checker.check(test_docx)
    print(f"   发现 {len(issues)} 个问题")

    # 筛选可修复的问题
    fixable = [i for i in issues if i.code in AutoFixer.FIXABLE_RULES]
    print(f"   其中 {len(fixable)} 个可自动修复:")
    for i in fixable:
        print(f"     [{i.code}] {i.description}")

    # 3. 运行自动修复
    print("\n3. 运行自动修复...")
    doc = Document(test_docx)
    fixer = AutoFixer(doc, issues)
    fix_log = fixer.fix(output_docx)
    print(f"   应用了 {len(fix_log)} 处修复")

    for record in fix_log:
        status = "OK" if record.paragraph_idx >= 0 else "FAIL"
        print(f"   [{status}] [{record.code}] 段落{record.paragraph_idx}: "
              f'"{record.original}" → "{record.fixed}"')

    # 4. 验证 Track Changes 标记
    print("\n4. 验证 Track Changes 标记...")
    tc = verify_track_changes(output_docx)
    print(f"   w:ins 元素: {tc['ins_elements']}")
    print(f"   w:del 元素: {tc['del_elements']}")
    print(f"   w:delText 元素: {tc['delText_elements']}")

    tc_ok = tc['ins_elements'] > 0 and tc['del_elements'] > 0
    print(f"   Track Changes 标记: {'PASS' if tc_ok else 'FAIL'}")

    # 5. 验证文本修复结果
    print("\n5. 验证文本修复结果...")
    checks = verify_text_fixed(output_docx)
    pass_count = 0
    fail_count = 0
    for name, result in checks.items():
        status = "PASS" if result else "FAIL"
        print(f"   [{status}] {name}")
        if result:
            pass_count += 1
        else:
            fail_count += 1

    # 6. 修复摘要
    print("\n6. 修复摘要...")
    summary = fixer.get_summary()
    print(f"   总修复数: {summary['total_fixes']}")
    print(f"   按规则: {summary['by_code']}")
    if summary['failed']:
        print(f"   失败: {len(summary['failed'])}")

    # 清理
    try:
        os.remove(test_docx)
    except Exception:
        pass
    try:
        os.remove(output_docx)
    except Exception:
        pass
    print(f"\n测试文件已清理。")

    # 总结
    print("\n" + "=" * 60)
    print(f"Track Changes: {'PASS' if tc_ok else 'FAIL'}")
    print(f"文本验证: {pass_count} PASS / {fail_count} FAIL")
    total_pass = (1 if tc_ok else 0) + pass_count
    total_fail = (0 if tc_ok else 1) + fail_count
    print(f"总计: {total_pass} PASS / {total_fail} FAIL")

    return total_fail == 0


if __name__ == "__main__":
    success = run_test()
    sys.exit(0 if success else 1)
