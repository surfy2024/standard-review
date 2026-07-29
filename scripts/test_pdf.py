#!/usr/bin/env python3
"""
PDF 支持测试脚本

1. 创建带已知问题的样例 PDF
2. 运行检查器验证规则检出
3. 创建相同样例的 .docx 对比结果一致性
"""

import os
import sys
import re
import tempfile

# 添加脚本目录到路径
script_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, script_dir)

from check_standard import StandardChecker, Severity


def create_test_pdf(pdf_path: str):
    """创建带已知问题的样例 PDF"""
    import fitz

    doc = fitz.open()  # 新建空文档
    page = doc.new_page(width=595, height=842)  # A4

    # 用 insert_text 逐行写入，模拟标准文档结构
    # 设置不同字体大小来模拟标题和正文

    y = 60
    line_height = 20

    def add_line(text, size=11, bold=False, font="helv"):
        nonlocal y
        # 中文字体用 china-s 或 china-t
        if re.search(r'[\u4e00-\u9fff]', text):
            fontname = "china-s"
        else:
            fontname = font
        page.insert_text((72, y), text, fontsize=size, fontname=fontname)
        y += line_height + (size - 11) * 1.5

    def add_heading(text, level=1):
        nonlocal y
        size = {1: 16, 2: 14, 3: 12}.get(level, 11)
        y += 10  # 标题前间距
        page.insert_text((72, y), text, fontsize=size, fontname="china-s")
        y += size * 1.8

    # === 目次 ===
    add_heading("目次", 1)
    add_line("前言..............I")
    add_line("1 范围............1")
    add_line("2 规范性引用文件..2")

    # === 前言 ===
    add_heading("前言", 1)
    add_line("本文件按照GB/T 1.1-2020的规定起草。")
    add_line("本文件必须遵循相关法律法规。")  # T001: 必须→应

    # === 范围 ===
    add_heading("1 范围", 1)
    add_line("本文件规定了,产品的技术要求和试验方法。")  # W001: 半角逗号
    add_line("长度为80mm，质量为50kg。")  # F002: 数字单位无空格

    # === 规范性引用文件（缺失，触发 S007）===
    # 故意不写

    # === 术语和定义 ===
    add_heading("3 术语和定义", 1)
    add_line("分辨率是指仪器能区分的最小距离。")  # W003: 术语未加粗

    # === 技术要求 ===
    add_heading("4 技术要求", 1)
    add_heading("4.1 外观要求。", 2)  # F001: 标题末尾标点
    add_line("产品应遵守以下要求。")  # T004: 遵守+要求
    add_line("外形尺寸为80x25x50 mm。")  # F003: 尺寸表述不规范
    add_line("由公式3计算得出。")  # F011: 公式编号无括号

    # === 孤立条 ===
    add_heading("5 试验方法", 1)
    add_heading("5.1 总则", 2)  # S004: 只有 5.1 无 5.2

    # === 引言含"应" ===
    add_heading("引言", 1)
    add_line("本标准的制定应满足行业需求。")  # S009: 引言含"应"

    # === 附录使用 I ===
    add_heading("附录 I", 1)  # S006: 附录编号用了 I

    doc.save(pdf_path)
    doc.close()
    print(f"测试 PDF 已生成：{pdf_path}")


def create_test_docx(docx_path: str):
    """创建与 PDF 内容相同的 .docx 用于对比"""
    from docx import Document

    doc = Document()

    # 目次
    doc.add_heading("目次", level=1)
    doc.add_paragraph("前言..............I")
    doc.add_paragraph("1 范围............1")
    doc.add_paragraph("2 规范性引用文件..2")

    # 前言
    doc.add_heading("前言", level=1)
    doc.add_paragraph("本文件按照GB/T 1.1-2020的规定起草。")
    doc.add_paragraph("本文件必须遵循相关法律法规。")  # T001

    # 范围
    doc.add_heading("1 范围", level=1)
    doc.add_paragraph("本文件规定了,产品的技术要求和试验方法。")  # W001
    doc.add_paragraph("长度为80mm，质量为50kg。")  # F002

    # 术语和定义
    doc.add_heading("3 术语和定义", level=1)
    doc.add_paragraph("分辨率是指仪器能区分的最小距离。")  # W003

    # 技术要求
    doc.add_heading("4 技术要求", level=1)
    doc.add_heading("4.1 外观要求。", level=2)  # F001
    doc.add_paragraph("产品应遵守以下要求。")  # T004
    doc.add_paragraph("外形尺寸为80x25x50 mm。")  # F003
    doc.add_paragraph("由公式3计算得出。")  # F011

    # 孤立条
    doc.add_heading("5 试验方法", level=1)
    doc.add_heading("5.1 总则", level=2)  # S004

    # 引言含"应"
    doc.add_heading("引言", level=1)
    doc.add_paragraph("本标准的制定应满足行业需求。")  # S009

    # 附录使用 I
    doc.add_heading("附录 I", level=1)  # S006

    doc.save(docx_path)
    print(f"测试 DOCX 已生成：{docx_path}")


def run_tests():
    """运行全部测试"""
    print("=" * 60)
    print("PDF 支持测试")
    print("=" * 60)

    # 创建临时文件
    tmp_dir = tempfile.mkdtemp()
    pdf_path = os.path.join(tmp_dir, "test_standard.pdf")
    docx_path = os.path.join(tmp_dir, "test_standard.docx")

    # 1. 创建测试文件
    print("\n1. 创建测试文件...")
    create_test_pdf(pdf_path)
    create_test_docx(docx_path)

    # 2. 检查 PDF
    print("\n2. 检查 PDF 文件...")
    pdf_checker = StandardChecker()
    pdf_issues = pdf_checker.check(pdf_path)
    pdf_codes = [i.code for i in pdf_issues]
    print(f"   PDF 检出 {len(pdf_issues)} 个问题：{sorted(set(pdf_codes))}")

    # 3. 检查 DOCX
    print("\n3. 检查 DOCX 文件...")
    docx_checker = StandardChecker()
    docx_issues = docx_checker.check(docx_path)
    docx_codes = [i.code for i in docx_issues]
    print(f"   DOCX 检出 {len(docx_issues)} 个问题：{sorted(set(docx_codes))}")

    # 4. 验证 PDF 检出了预期的规则
    print("\n4. 验证 PDF 规则检出...")
    expected_rules = {
        "T001": "禁用能愿动词'必须'",
        "W001": "半角标点",
        "F002": "数字单位无空格",
        "F001": "标题末尾标点",
        "T004": "遵守+要求搭配",
        "F003": "尺寸表述不规范",
        "F011": "公式编号无括号",
        "S004": "孤立条编号",
        "S009": "引言含'应'",
        "S006": "附录编号I",
        "S007": "无内容要素未声明",
    }

    pdf_pass = 0
    pdf_fail = 0
    for code, desc in expected_rules.items():
        found = code in pdf_codes
        status = "PASS" if found else "FAIL"
        if found:
            pdf_pass += 1
        else:
            pdf_fail += 1
        print(f"   [{status}] {code} - {desc}")

    # 5. 对比 PDF 和 DOCX 的检出一致性
    print("\n5. 对比 PDF 和 DOCX 检出一致性...")
    pdf_set = set(pdf_codes)
    docx_set = set(docx_codes)

    common = pdf_set & docx_set
    pdf_only = pdf_set - docx_set
    docx_only = docx_set - pdf_set

    print(f"   共同检出：{sorted(common)}")
    if pdf_only:
        print(f"   仅 PDF 检出：{sorted(pdf_only)}")
    if docx_only:
        print(f"   仅 DOCX 检出：{sorted(docx_only)}")

    consistency_pass = len(pdf_only) == 0 and len(docx_only) == 0
    if consistency_pass:
        print("   [PASS] PDF 和 DOCX 检出完全一致")
    else:
        # 允许 W003 和 T007 的差异（PDF bold 检测可能与 docx 不同）
        # W003 依赖 run.bold，PDF 的 bold 检测取决于字体
        # T007 依赖 Word 脚注 XML，PDF 无此结构
        allowed_diff = {"W003", "T007"}
        real_diff = (pdf_only | docx_only) - allowed_diff
        if not real_diff:
            print(f"   [PASS] 差异仅在允许范围内（W003/T007）")
            consistency_pass = True
        else:
            print(f"   [FAIL] 存在超出允许范围的差异：{sorted(real_diff)}")

    # 6. 验证 PDF 提取的段落结构
    print("\n6. 验证 PDF 段落结构...")
    from pdf_extractor import extract
    pdf_doc = extract(pdf_path)
    para_count = len(pdf_doc.paragraphs)
    heading_count = sum(1 for p in pdf_doc.paragraphs
                       if "Heading" in (p.style.name if p.style else ""))
    print(f"   段落总数：{para_count}")
    print(f"   标题数量：{heading_count}")

    struct_pass = para_count > 5 and heading_count > 3
    print(f"   [{'PASS' if struct_pass else 'FAIL'}] 段落结构提取")

    # 7. 测试纯文本提取
    print("\n7. 测试纯文本提取...")
    from pdf_extractor import extract_text
    raw_text = extract_text(pdf_path)
    text_pass = "范围" in raw_text and "前言" in raw_text and "术语" in raw_text
    print(f"   [{'PASS' if text_pass else 'FAIL'}] 纯文本提取包含关键内容")

    # 8. 测试错误输入处理
    print("\n8. 测试错误输入处理...")
    error_pass = True
    try:
        c = StandardChecker()
        c.check("nonexistent.pdf")
        error_pass = False  # 应该报错
    except SystemExit:
        pass  # 预期行为
    except Exception:
        error_pass = False
    print(f"   [{'PASS' if error_pass else 'FAIL'}] 不存在的文件处理")

    # 清理
    import shutil
    shutil.rmtree(tmp_dir, ignore_errors=True)

    # 汇总
    total_pass = pdf_pass + (1 if consistency_pass else 0) + (1 if struct_pass else 0) + (1 if text_pass else 0) + (1 if error_pass else 0)
    total = len(expected_rules) + 4
    print(f"\n{'=' * 60}")
    print(f"总计：{total_pass}/{total} 通过")
    print(f"  - PDF 规则检出：{pdf_pass}/{len(expected_rules)}")
    print(f"  - PDF/DOCX 一致性：{'PASS' if consistency_pass else 'FAIL'}")
    print(f"  - 段落结构提取：{'PASS' if struct_pass else 'FAIL'}")
    print(f"  - 纯文本提取：{'PASS' if text_pass else 'FAIL'}")
    print(f"  - 错误输入处理：{'PASS' if error_pass else 'FAIL'}")
    print(f"{'=' * 60}")

    return total_pass == total


if __name__ == "__main__":
    success = run_tests()
    sys.exit(0 if success else 1)
