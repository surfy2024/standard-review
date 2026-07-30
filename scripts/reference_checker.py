#!/usr/bin/env python3
"""
引用标准符合性检查模块

功能：
1. 自动层 (R001-R008): 引用标准提取、格式校验、交叉引用一致性
2. 用户提交层 (R009-R012): 要求性条款提取、语义匹配、指标级比对

使用方式：
    由 check_standard.py 调用，无需单独运行
"""

import re
import os
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, field
from enum import Enum


# ============================================================
# 数据结构
# ============================================================

@dataclass
class StandardReference:
    """引用标准条目"""
    raw_text: str           # 原始文本
    number: str             # 标准编号（如 GB/T 1.1-2020）
    prefix: str             # 前缀（如 GB/T, GB, YY/T）
    number_part: str        # 数字部分（如 1.1）
    year: str               # 年代号（如 2020），可能为空
    title: str              # 标准名称（可能为空）
    location: str           # 位置描述
    is_dated: bool          # 是否注日期引用


@dataclass
class RequirementClause:
    """要求性条款"""
    text: str               # 条款全文
    modal_verb: str         # 能愿动词（应/必须/不得/宜/可/不应）
    location: str           # 位置描述
    keywords: List[str]     # 关键词（用于匹配）
    indicators: List[Dict]  # 指标（数值+单位+比较符）
    source: str             # 来源（draft 或 reference + 文件名）


@dataclass
class ComplianceResult:
    """符合性比对结果"""
    draft_req: RequirementClause
    ref_req: RequirementClause
    status: str             # "不符合" / "符合" / "优于" / "无法判断"
    detail: str             # 比对详情
    similarity: float       # 相似度


# ============================================================
# 正则模式
# ============================================================

# 标准编号通用模式
# 覆盖：GB/T, GB, YY/T, YY, JB/T, HJ, DB11/T, Q/ABC, T/CAS, ISO, IEC 等
STANDARD_NUMBER_RE = re.compile(
    r'(?P<prefix>'
    r'GB/?T|GB|'           # 国家标准
    r'[A-Z]{2}/?T|'        # 行业推荐性标准 (YY/T, JB/T, HJ/T ...)
    r'[A-Z]{2}|'           # 行业强制性标准 (YY, JB, HJ ...)
    r'DB\d{2}/?T|'         # 地方推荐性标准
    r'DB\d{2}|'            # 地方强制性标准
    r'Q/[A-Z]+|'           # 企业标准
    r'T/[A-Z]+|'           # 团体标准
    r'ISO|'                # 国际标准
    r'IEC'                 # 国际电工标准
    r')'
    r'\s*'
    r'(?P<number>\d+\.?\d*)'
    r'(?:\s*[-—:：]\s*(?P<year>\d{4}))?'
)

# 能愿动词模式（用于识别要求性条款）
MODAL_VERBS = ['不得', '不应', '必须', '应', '宜', '可']
MODAL_VERB_RE = re.compile(r'(不得|不应|必须|应|宜|可)')

# 标准引用引用语模式（正文中引用标准时的常见表述）
CITATION_PATTERNS = [
    re.compile(r'按\s*(?:照\s*)?(GB/?T\s*\d+\.?\d*(?:\s*[-—:：]\s*\d{4})?)\s*(?:的\s*)?规定'),
    re.compile(r'符合\s*(GB/?T\s*\d+\.?\d*(?:\s*[-—:：]\s*\d{4})?)'),
    re.compile(r'见\s*(GB/?T\s*\d+\.?\d*(?:\s*[-—:：]\s*\d{4})?)'),
    re.compile(r'按照\s*(GB/?T\s*\d+\.?\d*(?:\s*[-—:：]\s*\d{4})?)'),
    re.compile(r'依据\s*(GB/?T\s*\d+\.?\d*(?:\s*[-—:：]\s*\d{4})?)'),
    re.compile(r'参照\s*(GB/?T\s*\d+\.?\d*(?:\s*[-—:：]\s*\d{4})?)'),
    # 通用模式：标准编号直接出现在文本中
    re.compile(r'(GB/?T\s*\d+\.?\d*(?:\s*[-—:：]\s*\d{4})?)'),
    re.compile(r'(YY/?T\s*\d+\.?\d*(?:\s*[-—:：]\s*\d{4})?)'),
    re.compile(r'(HJ/?T?\s*\d+\.?\d*(?:\s*[-—:：]\s*\d{4})?)'),
    re.compile(r'(JB/?T\s*\d+\.?\d*(?:\s*[-—:：]\s*\d{4})?)'),
    re.compile(r'(DB\d{2}/?T?\s*\d+(?:\s*[-—:：]\s*\d{4})?)'),
    re.compile(r'(ISO\s+\d+(?:\s*[-—:：]\s*\d{4})?)'),
    re.compile(r'(IEC\s+\d+(?:\s*[-—:：]\s*\d{4})?)'),
]

# 指标提取模式（数值 + 比较符 + 单位）
INDICATOR_RE = re.compile(
    r'(?P<comp>不低于|不大于|不超过|不少于|大于等于|小于等于|大于|小于|≥|≤|>|<|=|至少|最多|最小|最大)'
    r'\s*'
    r'(?P<value>\d+\.?\d*)'
    r'\s*'
    r'(?P<unit>%|‰|mg/kg|mg/mL|mg/L|g/kg|g/L|μg/L|ppm|mm|cm|km|GHz|MHz|Hz|kV|V|A|W|kW|MW|℃|°C|dB|lux|lx|次|个|条|件|人|天|日|小时|分钟|秒|年|月|周|类|种|项)?'
)

# 反向指标模式（数值在比较符前面）
INDICATOR_REVERSE_RE = re.compile(
    r'(?P<value>\d+\.?\d*)'
    r'\s*'
    r'(?P<unit>%|‰|mg/kg|mg/mL|mg/L|g/kg|g/L|μg/L|ppm|℃|°C|dB|Hz|MHz|GHz|kV|V|A|W|kW|MW|lux|lx)'
    r'\s*'
    r'(?P<op>以上|以下|以内|以外|及以上|及以下)'
)

# 规范性引用文件引导语
INTRO_TEXT_RE = re.compile(
    r'下列文件中的内容通过文中的规范性引用而构成本文件'
    r'(?:必不可少的条款|必不可少的条款|必不可少的组成条款)'
    r'。其中，注(?:日期|年代号)的引用文件，?仅该日期对应的版本适用于本文件；?'
    r'未注(?:日期|年代号)的引用文件，?其最新版本（包括所有的修改单）适用于本文件。'
)

# 宽松匹配引导语（允许部分差异）
INTRO_TEXT_LOOSE_RE = re.compile(
    r'下列文件.*?通过.*?规范性引用.*?构成本文件.*?条款'
)


# ============================================================
# 辅助函数
# ============================================================

def _normalize_std_number(num: str) -> str:
    """标准化标准编号，用于比较"""
    # 移除所有空格
    num = re.sub(r'\s+', '', num)
    # 统一连接符为一字线
    num = num.replace('—', '-').replace('–', '-').replace('：', ':')
    # 统一冒号为一字线（GB/T 1.1:2020 → GB/T 1.1-2020）
    if ':' in num and re.search(r':\d{4}$', num):
        num = re.sub(r':(\d{4})$', r'-\1', num)
    return num


def _extract_std_prefix(num: str) -> str:
    """提取标准编号前缀"""
    match = re.match(r'^([A-Za-z/]+)', num.strip())
    if match:
        return match.group(1).upper()
    return ""


def _split_sentences(text: str) -> List[str]:
    """将文本分割为句子"""
    # 按句号、分号、换行分割
    sentences = re.split(r'[。；;\n！!？?]+', text)
    return [s.strip() for s in sentences if s.strip() and len(s.strip()) > 3]


def _extract_keywords(text: str) -> List[str]:
    """从文本中提取关键词（简化版）"""
    # 移除常见停用词和标点
    stop_words = {'的', '了', '和', '与', '或', '及', '在', '为', '对', '按', '按照',
                  '根据', '依据', '参照', '符合', '见', '本', '该', '其', '应', '必须',
                  '不得', '宜', '可', '不应', '不', '等', '中', '上', '下', '内', '外',
                  '以', '是', '有', '无', '一个', '一种', '这个', '那个'}
    # 提取中文词语（2-6字）和英文单词
    cn_words = re.findall(r'[\u4e00-\u9fff]{2,6}', text)
    en_words = re.findall(r'[A-Za-z]{3,}', text)
    
    keywords = []
    for w in cn_words:
        if w not in stop_words:
            keywords.append(w)
    for w in en_words:
        if w.lower() not in stop_words:
            keywords.append(w.lower())
    
    return keywords


def _extract_indicators(text: str) -> List[Dict]:
    """从文本中提取指标（数值+比较符+单位）"""
    indicators = []
    
    # 正向指标：不低于 20, ≥ 0.5, 大于 100
    for match in INDICATOR_RE.finditer(text):
        comp = match.group('comp')
        value = float(match.group('value'))
        unit = match.group('unit') or ''
        indicators.append({
            'comp': comp,
            'value': value,
            'unit': unit,
            'raw': match.group(),
            'direction': _get_comp_direction(comp),
        })
    
    # 反向指标：20% 以上, 0.5mg/L 以下
    for match in INDICATOR_REVERSE_RE.finditer(text):
        value = float(match.group('value'))
        unit = match.group('unit')
        op = match.group('op')
        direction = 'min' if '以上' in op or '及以上' in op else 'max'
        comp = '不低于' if direction == 'min' else '不大于'
        indicators.append({
            'comp': comp,
            'value': value,
            'unit': unit,
            'raw': match.group(),
            'direction': direction,
        })
    
    return indicators


def _get_comp_direction(comp: str) -> str:
    """获取比较方向：min（最小值约束）或 max（最大值约束）"""
    min_comps = {'不低于', '大于等于', '≥', '>', '大于', '至少', '最小', '不少于'}
    max_comps = {'不大于', '不超过', '小于等于', '≤', '<', '小于', '最多', '最大'}
    if comp in min_comps:
        return 'min'
    if comp in max_comps:
        return 'max'
    return 'eq'  # 等于


def _compute_similarity(text1_keywords: List[str], text2_keywords: List[str]) -> float:
    """计算两个关键词列表的相似度（Jaccard 系数）"""
    if not text1_keywords or not text2_keywords:
        return 0.0
    set1 = set(text1_keywords)
    set2 = set(text2_keywords)
    intersection = set1 & set2
    union = set1 | set2
    return len(intersection) / len(union) if union else 0.0


def _load_document(file_path: str) -> List[Dict]:
    """加载文档（支持 .docx 和 .pdf），返回段落列表"""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在：{path}")
    
    ext = path.suffix.lower()
    paragraphs = []
    
    if ext == '.pdf':
        try:
            from pdf_extractor import extract as extract_pdf
        except ImportError:
            raise ImportError("需要安装 PyMuPDF 库来处理 PDF 文件")
        doc = extract_pdf(str(path))
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                paragraphs.append({"index": i, "text": text, "section": ""})
    elif ext == '.docx':
        try:
            from docx import Document
        except ImportError:
            raise ImportError("需要安装 python-docx 库来处理 DOCX 文件")
        doc = Document(str(path))
        current_section = ""
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style else ""
            is_heading = ("Heading" in style_name or "标题" in style_name)
            if is_heading:
                current_section = text
            paragraphs.append({"index": i, "text": text, "section": current_section})
    else:
        raise ValueError(f"不支持的文件格式：{ext}")
    
    return paragraphs


# ============================================================
# 主检查器
# ============================================================

class ReferenceChecker:
    """引用标准符合性检查器"""
    
    def __init__(self, paragraphs: List[Dict], headings: List[Dict],
                 section_map: Dict[int, str], add_issue_func=None):
        """
        Args:
            paragraphs: 段落列表
            headings: 标题列表
            section_map: 段落索引→章节名映射
            add_issue_func: 添加问题的回调函数
        """
        self.paragraphs = paragraphs
        self.headings = headings
        self.section_map = section_map
        self._add_issue = add_issue_func
        
        # 提取结果
        self.references: List[StandardReference] = []
        self.citations: List[Dict] = []  # 正文中的引用
        self.ref_section_paras: List[Dict] = []  # 规范性引用文件章节的段落
        
        # 符合性检查结果
        self.draft_requirements: List[RequirementClause] = []
        self.ref_requirements: List[RequirementClause] = []
        self.compliance_results: List[ComplianceResult] = []
        self.matched_pairs_count: int = 0  # R011 匹配到的要求对数
    
    # ============================================================
    # 自动层 R001-R008
    # ============================================================
    
    def check_auto(self) -> None:
        """执行自动层检查（无需用户提供额外文件）"""
        self._extract_reference_section()
        self._extract_references()       # R001
        self._check_reference_format()    # R002
        self._check_duplicate_refs()      # R003
        self._check_reference_order()     # R004
        self._check_intro_text()          # R005
        self._extract_citations()         # R006
        self._check_missing_refs()        # R007
        self._check_redundant_refs()      # R008

    def get_reference_numbers(self) -> List[str]:
        """获取引用标准号列表（供自动下载使用）

        Returns:
            标准号列表，如 ['GB/T 1.1-2020', 'GB 3836.1-2021', ...]
        """
        return [ref.number for ref in self.references if ref.number]
    
    def _extract_reference_section(self) -> None:
        """提取"规范性引用文件"章节的段落"""
        in_ref_section = False
        ref_heading_index = -1
        
        for i, para in enumerate(self.paragraphs):
            text = para["text"].strip()
            
            if in_ref_section:
                # 已在引用文件章节内：检查是否遇到下一个标题（章节结束）
                if not self._is_toc_entry(para) and (para.get("is_heading") or self._is_heading_text(text)):
                    if "规范性引用文件" not in text.replace(" ", "").replace("　", ""):
                        in_ref_section = False
                        break
                # 收集引用文件章节内的所有段落（包括样式异常的条目）
                self.ref_section_paras.append(para)
            else:
                # 尚未进入引用文件章节：跳过目次条目，避免误识别为标题
                if self._is_toc_entry(para):
                    continue
                if para.get("is_heading") or self._is_heading_text(text):
                    if "规范性引用文件" in text.replace(" ", "").replace("　", ""):
                        in_ref_section = True
                        ref_heading_index = i
    
    def _is_toc_entry(self, para: Dict) -> bool:
        """判断段落是否为目次（TOC）条目"""
        # 方法1：样式名包含 "toc"（不区分大小写）
        style = para.get("style", "")
        if style and "toc" in style.lower():
            return True
        # 方法2：文本含制表符后跟页码（如 "2 规范性引用文件\t1"）
        text = para["text"].strip()
        if re.search(r'\t\s*\d+\s*$', text):
            return True
        return False
    
    def _is_heading_text(self, text: str) -> bool:
        """判断文本是否为标题"""
        text = text.strip()
        if not text or len(text) > 50:
            return False
        # 排除目次条目（含制表符 + 页码）
        if re.search(r'\t\s*\d+\s*$', text):
            return False
        # 编号 + 标题：如 "2 规范性引用文件"
        if re.match(r'^\d+\s+\S', text):
            return True
        # 已知的章节名
        known_headings = ["规范性引用文件", "术语和定义", "术语", "定义", "技术要求",
                          "试验方法", "检验规则", "标志、包装、运输和贮存", "范围",
                          "前言", "引言", "附录"]
        for h in known_headings:
            if h in text:
                return True
        return False
    
    def _extract_references(self) -> None:
        """R001: 从规范性引用文件章节提取引用标准"""
        for para in self.ref_section_paras:
            text = para["text"].strip()
            if not text:
                continue
            
            # 跳过引导语
            if INTRO_TEXT_LOOSE_RE.search(text):
                continue
            
            # 尝试匹配标准编号
            matches = list(STANDARD_NUMBER_RE.finditer(text))
            if matches:
                for match in matches:
                    prefix = match.group('prefix')
                    number_part = match.group('number')
                    year = match.group('year') or ''
                    
                    # 构建标准编号
                    std_number = f"{prefix} {number_part}"
                    if year:
                        std_number += f"-{year}"
                    
                    # 提取标题（编号后面的部分）
                    title_start = match.end()
                    title = text[title_start:].strip()
                    # 清理标题
                    title = re.sub(r'^[\s,，。、:：·]+', '', title)
                    # 截断到合理长度
                    if len(title) > 100:
                        title = title[:100] + "..."
                    
                    self.references.append(StandardReference(
                        raw_text=text,
                        number=std_number,
                        prefix=prefix.upper(),
                        number_part=number_part,
                        year=year,
                        title=title,
                        location=f"规范性引用文件, 段落 {para['index']}",
                        is_dated=bool(year),
                    ))
    
    def _check_reference_format(self) -> None:
        """R002: 校验引用标准格式"""
        for ref in self.references:
            # 检查编号格式：应有空格分隔前缀和数字
            if not re.match(r'^[A-Za-z/]+\s+\d', ref.number):
                self._safe_add_issue(
                    "R002", "WARNING",
                    ref.location,
                    f"引用标准编号格式不规范：{ref.number}",
                    "前缀与数字之间应有一个空格，如 GB/T 1.1-2020",
                    ref.raw_text
                )
            
            # 检查注日期引用的连接符（应为一字线 -）
            if ref.year:
                # 检查是否使用了冒号或全角连接符
                if '：' in ref.number or ':' in ref.number:
                    self._safe_add_issue(
                        "R002", "ERROR",
                        ref.location,
                        f"注日期引用使用了冒号连接：{ref.number}",
                        "应使用一字线连接，如 GB/T 1.1-2020",
                        ref.raw_text
                    )
                if '—' in ref.number or '–' in ref.number:
                    self._safe_add_issue(
                        "R002", "WARNING",
                        ref.location,
                        f"注日期引用使用了全角连接符：{ref.number}",
                        "应使用半角一字线（-）连接",
                        ref.raw_text
                    )
            
            # 检查前缀格式
            prefix_upper = ref.prefix.upper()
            if prefix_upper.startswith('GB') and '/T' not in prefix_upper and 'GBT' not in prefix_upper:
                # GB 而非 GB/T，可能是强制性国家标准，检查是否有注明
                pass  # 强制性标准不需要 /T，这是合法的
            
            # 检查标准编号中是否有多余空格
            if re.search(r'\s{2,}', ref.number):
                self._safe_add_issue(
                    "R002", "SUGGESTION",
                    ref.location,
                    f"引用标准编号含有多余空格：{ref.number}",
                    "编号中应仅有一个空格分隔前缀和数字",
                    ref.raw_text
                )
    
    def _check_duplicate_refs(self) -> None:
        """R003: 检查重复引用"""
        seen = {}
        for ref in self.references:
            norm = _normalize_std_number(ref.number)
            if norm in seen:
                self._safe_add_issue(
                    "R003", "WARNING",
                    ref.location,
                    f"引用标准重复：{ref.number}",
                    f"该标准已在 {seen[norm].location} 引用，请删除重复项",
                    ref.raw_text
                )
            else:
                seen[norm] = ref
    
    def _check_reference_order(self) -> None:
        """R004: 检查引用标准排序规范性"""
        if len(self.references) < 2:
            return
        
        # GB/T 1.1-2020 要求：先国家、再行业、再地方、再国际
        # 同类内按编号升序排列
        type_order = {
            'GB': 0, 'GB/T': 0, 'GBT': 0,
            'GBN': 0,
        }
        # 行业标准前缀
        industry_prefixes = ['YY', 'YY/T', 'JB', 'JB/T', 'HJ', 'HJ/T', 'QB', 'QB/T',
                            'SJ', 'SJ/T', 'HG', 'HG/T', 'SY', 'SY/T', 'DL', 'DL/T',
                            'JG', 'JG/T', 'CJ', 'CJ/T', 'GA', 'GA/T', 'WS', 'WS/T',
                            'SL', 'SL/T', 'MT', 'MT/T', 'YS', 'YS/T', 'LD', 'MZ',
                            'YY', 'SH', 'SH/T', 'NB', 'NB/T', 'JT', 'JT/T']
        for p in industry_prefixes:
            type_order[p] = 1
        
        local_prefixes = ['DB11', 'DB12', 'DB13', 'DB14', 'DB21', 'DB22', 'DB23',
                         'DB31', 'DB32', 'DB33', 'DB34', 'DB35', 'DB36', 'DB37',
                         'DB41', 'DB42', 'DB43', 'DB44', 'DB45', 'DB46', 'DB50',
                         'DB51', 'DB52', 'DB53', 'DB54', 'DB61', 'DB62', 'DB63',
                         'DB64', 'DB65']
        for p in local_prefixes:
            type_order[p] = 2
            type_order[p + '/T'] = 2
        
        type_order['ISO'] = 3
        type_order['IEC'] = 3
        
        # 企业/团体标准
        for ref in self.references:
            if ref.prefix.startswith('Q/'):
                type_order[ref.prefix] = 4
            elif ref.prefix.startswith('T/'):
                type_order[ref.prefix] = 5
        
        # 检查排序
        prev_type_order = -1
        prev_number = ""
        for i, ref in enumerate(self.references):
            current_type_order = type_order.get(ref.prefix, 99)
            current_number = ref.number_part
            
            # 类型顺序检查
            if current_type_order < prev_type_order:
                self._safe_add_issue(
                    "R004", "SUGGESTION",
                    ref.location,
                    f"引用标准排序不规范：{ref.number}",
                    "引用标准应按国家标准→行业标准→地方标准→国际标准的顺序排列",
                    ""
                )
            
            # 同类型内编号顺序检查
            if current_type_order == prev_type_order and current_number < prev_number:
                self._safe_add_issue(
                    "R004", "SUGGESTION",
                    ref.location,
                    f"同类型引用标准编号未按升序排列：{ref.number}",
                    f"建议按编号升序排列（{ref.number} 应排在前面）",
                    ""
                )
            
            prev_type_order = current_type_order
            prev_number = current_number
    
    def _check_intro_text(self) -> None:
        """R005: 检查规范性引用文件引导语"""
        if not self.ref_section_paras:
            return
        
        # 查找引导语
        found_intro = False
        for para in self.ref_section_paras:
            text = para["text"].strip()
            if INTRO_TEXT_LOOSE_RE.search(text):
                found_intro = True
                # 检查引导语是否完整
                if not INTRO_TEXT_RE.search(text):
                    # 引导语不完整
                    self._safe_add_issue(
                        "R005", "WARNING",
                        f"规范性引用文件, 段落 {para['index']}",
                        "规范性引用文件引导语不完整",
                        "引导语应包含：注日期引用仅该日期版本适用；未注日期引用其最新版本适用",
                        text[:200]
                    )
                break
        
        if not found_intro and self.references:
            self._safe_add_issue(
                "R005", "ERROR",
                "规范性引用文件",
                "缺少规范性引用文件引导语",
                "应在引用文件列表前添加引导语：\"下列文件中的内容通过文中的规范性引用而构成本文件必不可少的条款...\"",
                ""
            )
    
    def _extract_citations(self) -> None:
        """R006: 从全文提取标准引用"""
        # 确定规范性引用文件章节的范围
        ref_para_indices = set(p["index"] for p in self.ref_section_paras)
        
        for para in self.paragraphs:
            if para["index"] in ref_para_indices:
                continue  # 跳过规范性引用文件章节本身
            
            text = para["text"]
            for pattern in CITATION_PATTERNS:
                for match in pattern.finditer(text):
                    cited_num = match.group(1) if match.groups() else match.group()
                    self.citations.append({
                        "number": cited_num.strip(),
                        "normalized": _normalize_std_number(cited_num.strip()),
                        "location": f"段落 {para['index']}",
                        "context": text[:200],
                    })
    
    def _check_missing_refs(self) -> None:
        """R007: 检查正文中引用但未在规范性引用文件中列出的标准"""
        ref_numbers = set(_normalize_std_number(r.number) for r in self.references)
        
        # 去重
        seen_missing = set()
        for citation in self.citations:
            norm = citation["normalized"]
            if norm not in ref_numbers and norm not in seen_missing:
                seen_missing.add(norm)
                self._safe_add_issue(
                    "R007", "ERROR",
                    citation["location"],
                    f"正文中引用了标准 {citation['number']}，但未在\"规范性引用文件\"中列出",
                    f"请在\"规范性引用文件\"中补充 {citation['number']} 的引用",
                    citation["context"]
                )
    
    def _check_redundant_refs(self) -> None:
        """R008: 检查规范性引用文件中列出但正文从未引用的标准"""
        citation_numbers = set(c["normalized"] for c in self.citations)
        
        for ref in self.references:
            norm = _normalize_std_number(ref.number)
            # 检查正文是否有引用（也检查不带年份的引用）
            norm_no_year = re.sub(r'[-—:：]\d{4}$', '', norm)
            found = norm in citation_numbers or norm_no_year in citation_numbers
            
            # 也检查引用列表中是否有不带年份的版本
            if not found:
                for c in citation_numbers:
                    c_no_year = re.sub(r'[-—:：]\d{4}$', '', c)
                    if c_no_year == norm_no_year:
                        found = True
                        break
            
            if not found:
                self._safe_add_issue(
                    "R008", "WARNING",
                    ref.location,
                    f"引用标准 {ref.number} 在正文中未被引用",
                    "规范性引用文件中列出的标准应在正文中至少引用一次，或从引用文件中删除",
                    ref.raw_text
                )
    
    # ============================================================
    # 用户提交层 R009-R012
    # ============================================================
    
    def check_compliance(self, ref_files: List[str]) -> None:
        """执行用户提交层检查（需要用户提供引用标准文档）
        
        Args:
            ref_files: 引用标准文档路径列表
        """
        print("\n=== 引用标准符合性检查 ===")
        
        # R009: 从引用标准文档中提取要求性条款
        self.ref_requirements = self._extract_requirements_from_files(ref_files)
        print(f"从 {len(ref_files)} 个引用标准文档中提取了 {len(self.ref_requirements)} 条要求性条款")
        
        if not self.ref_requirements:
            print("警告：未从引用标准文档中提取到要求性条款，跳过符合性检查")
            return
        
        # R010: 从草稿中提取要求性条款
        self.draft_requirements = self._extract_requirements_from_draft()
        print(f"从草稿中提取了 {len(self.draft_requirements)} 条要求性条款")
        
        if not self.draft_requirements:
            print("警告：未从草稿中提取到要求性条款，跳过符合性检查")
            return
        
        # R011: 语义匹配
        matched_pairs = self._match_requirements()
        self.matched_pairs_count = len(matched_pairs)
        print(f"匹配到 {len(matched_pairs)} 对相关要求")
        
        # R012: 指标级比对
        self._compare_indicators(matched_pairs)
    
    def _extract_requirements_from_files(self, file_paths: List[str]) -> List[RequirementClause]:
        """R009: 从用户提供的引用标准文档中提取要求性条款"""
        requirements = []
        
        for file_path in file_paths:
            file_name = Path(file_path).name
            print(f"  正在提取要求性条款：{file_name}")
            
            try:
                paragraphs = _load_document(file_path)
            except Exception as e:
                print(f"  警告：无法加载文件 {file_name}：{e}")
                continue
            
            for para in paragraphs:
                text = para["text"].strip()
                if not text:
                    continue
                
                # 分句
                sentences = _split_sentences(text)
                for sent in sentences:
                    # 检查是否包含能愿动词
                    modal_match = MODAL_VERB_RE.search(sent)
                    if modal_match:
                        modal_verb = modal_match.group(1)
                        keywords = _extract_keywords(sent)
                        indicators = _extract_indicators(sent)
                        
                        requirements.append(RequirementClause(
                            text=sent,
                            modal_verb=modal_verb,
                            location=f"{file_name}, 段落 {para['index']}",
                            keywords=keywords,
                            indicators=indicators,
                            source=f"reference:{file_name}",
                        ))
            
            print(f"    → 提取到 {sum(1 for r in requirements if r.source == f'reference:{file_name}')} 条")
        
        return requirements
    
    def _extract_requirements_from_draft(self) -> List[RequirementClause]:
        """R010: 从草稿中提取要求性条款"""
        requirements = []
        
        # 跳过规范性引用文件章节（引用文件不是要求性条款）
        ref_para_indices = set(p["index"] for p in self.ref_section_paras)
        
        for para in self.paragraphs:
            if para["index"] in ref_para_indices:
                continue
            
            text = para["text"].strip()
            if not text:
                continue
            
            # 分句
            sentences = _split_sentences(text)
            for sent in sentences:
                modal_match = MODAL_VERB_RE.search(sent)
                if modal_match:
                    modal_verb = modal_match.group(1)
                    keywords = _extract_keywords(sent)
                    indicators = _extract_indicators(sent)
                    
                    requirements.append(RequirementClause(
                        text=sent,
                        modal_verb=modal_verb,
                        location=f"草稿, 段落 {para['index']}",
                        keywords=keywords,
                        indicators=indicators,
                        source="draft",
                    ))
        
        return requirements
    
    def _match_requirements(self) -> List[Tuple[RequirementClause, RequirementClause, float]]:
        """R011: 语义匹配——将草稿要求与引用标准要求按主题关联
        
        Returns:
            匹配对列表：(draft_req, ref_req, similarity)
        """
        matched_pairs = []
        used_ref_indices = set()
        
        SIMILARITY_THRESHOLD = 0.15  # 相似度阈值
        
        for draft_req in self.draft_requirements:
            best_match = None
            best_sim = 0.0
            
            for j, ref_req in enumerate(self.ref_requirements):
                if j in used_ref_indices:
                    continue
                
                sim = _compute_similarity(draft_req.keywords, ref_req.keywords)
                if sim > best_sim:
                    best_sim = sim
                    best_match = j
            
            if best_match is not None and best_sim >= SIMILARITY_THRESHOLD:
                matched_pairs.append((draft_req, self.ref_requirements[best_match], best_sim))
                used_ref_indices.add(best_match)
        
        return matched_pairs
    
    def _compare_indicators(self, matched_pairs: List[Tuple[RequirementClause, RequirementClause, float]]) -> None:
        """R012: 指标级比对——草稿要求 vs 引用标准要求"""

        for draft_req, ref_req, similarity in matched_pairs:
            # 如果两边都没有指标，无法比较，跳过
            if not draft_req.indicators and not ref_req.indicators:
                continue

            # 提取引用标准来源名称
            ref_source = self._extract_source_name(ref_req.source)

            # 如果草稿有指标但引用标准没有指标，无法判断
            if draft_req.indicators and not ref_req.indicators:
                context = (
                    f"【草稿条款】\n{draft_req.text}\n\n"
                    f"【引用标准条款】（来源：{ref_source}）\n{ref_req.text}\n\n"
                    f"【分析】草稿含定量指标但引用标准中未找到对应指标，无法判断符合性\n"
                    f"匹配相似度: {similarity:.0%}"
                )
                self._safe_add_issue(
                    "R012", "SUGGESTION",
                    draft_req.location,
                    f"草稿要求含定量指标，但引用标准（{ref_source}）中未找到对应指标，无法判断符合性",
                    f"建议人工核查草稿指标是否与{ref_source}的要求一致",
                    context
                )
                continue

            # 如果引用标准有指标但草稿没有
            if not draft_req.indicators and ref_req.indicators:
                ref_ind_str = "；".join(
                    f"{ind['comp']} {ind['value']}{ind['unit']}" for ind in ref_req.indicators
                )
                context = (
                    f"【草稿条款】\n{draft_req.text}\n\n"
                    f"【引用标准条款】（来源：{ref_source}）\n{ref_req.text}\n\n"
                    f"【引用标准指标】{ref_ind_str}\n\n"
                    f"【分析】引用标准有定量指标但草稿中未规定，可能导致不符合\n"
                    f"匹配相似度: {similarity:.0%}"
                )
                self._safe_add_issue(
                    "R012", "WARNING",
                    draft_req.location,
                    f"引用标准（{ref_source}）有定量指标但草稿中未规定，可能导致不符合",
                    f"建议在草稿中补充引用标准{ref_source}的指标要求：{ref_ind_str}",
                    context
                )
                continue

            # 两边都有指标，进行比对
            for d_ind in draft_req.indicators:
                for r_ind in ref_req.indicators:
                    # 只有单位相同（或一方无单位）才比较
                    if d_ind['unit'] and r_ind['unit'] and d_ind['unit'] != r_ind['unit']:
                        continue

                    result = self._compare_single_indicator(d_ind, r_ind)

                    d_ind_str = f"{d_ind['comp']} {d_ind['value']}{d_ind['unit']}"
                    r_ind_str = f"{r_ind['comp']} {r_ind['value']}{r_ind['unit']}"

                    if result == "不符合":
                        context = (
                            f"【草稿条款】\n{draft_req.text}\n\n"
                            f"【引用标准条款】（来源：{ref_source}）\n{ref_req.text}\n\n"
                            f"【指标对比】\n"
                            f"  草稿指标：{d_ind_str}\n"
                            f"  引用标准指标：{r_ind_str}\n\n"
                            f"【对比结果】草稿要求低于引用标准要求\n"
                            f"  根据GB/T 1.1-2020，新标准的要求不应低于现行标准的要求。\n"
                            f"匹配相似度: {similarity:.0%}"
                        )
                        self._safe_add_issue(
                            "R012", "ERROR",
                            draft_req.location,
                            f"草稿指标「{d_ind_str}」不符合引用标准{ref_source}的要求「{r_ind_str}」",
                            f"草稿要求低于引用标准要求，应调整草稿指标使其不低于{ref_source}的要求（{r_ind_str}）",
                            context
                        )
                    elif result == "优于":
                        context = (
                            f"【草稿条款】\n{draft_req.text}\n\n"
                            f"【引用标准条款】（来源：{ref_source}）\n{ref_req.text}\n\n"
                            f"【指标对比】\n"
                            f"  草稿指标：{d_ind_str}\n"
                            f"  引用标准指标：{r_ind_str}\n\n"
                            f"【对比结果】草稿要求严于引用标准，符合要求\n"
                            f"匹配相似度: {similarity:.0%}"
                        )
                        self._safe_add_issue(
                            "R012", "SUGGESTION",
                            draft_req.location,
                            f"草稿指标「{d_ind_str}」优于引用标准{ref_source}的要求「{r_ind_str}」",
                            f"草稿要求严于引用标准，符合要求，无需修改",
                            context
                        )
                    elif result == "符合":
                        context = (
                            f"【草稿条款】\n{draft_req.text}\n\n"
                            f"【引用标准条款】（来源：{ref_source}）\n{ref_req.text}\n\n"
                            f"【指标对比】\n"
                            f"  草稿指标：{d_ind_str}\n"
                            f"  引用标准指标：{r_ind_str}\n\n"
                            f"【对比结果】草稿要求与引用标准一致\n"
                            f"匹配相似度: {similarity:.0%}"
                        )
                        self._safe_add_issue(
                            "R012", "SUGGESTION",
                            draft_req.location,
                            f"草稿指标「{d_ind_str}」与引用标准{ref_source}的要求一致",
                            f"草稿要求与引用标准一致，符合要求",
                            context
                        )

    def _extract_source_name(self, source: str) -> str:
        """从 source 字段提取引用标准名称

        如 "reference:GB_3836.16-2024.pdf" → "GB 3836.16-2024"
        """
        if not source or source == "draft":
            return "草稿"
        # 格式: "reference:GB_3836.16-2024.pdf"
        name = source.replace("reference:", "").replace(".pdf", "")
        # 将下划线转为空格: "GB_3836.16-2024" → "GB 3836.16-2024"
        name = name.replace("_", " ")
        return name

    def _compare_single_indicator(self, draft_ind: Dict, ref_ind: Dict) -> str:
        """比较单个指标
        
        Returns:
            "不符合" / "符合" / "优于" / "无法判断"
        """
        d_dir = draft_ind['direction']
        r_dir = ref_ind['direction']
        d_val = draft_ind['value']
        r_val = ref_ind['value']
        
        # 同方向比较
        if d_dir == r_dir:
            if d_dir == 'min':
                # 最小值约束：草稿的最小值 >= 引用标准的最小值 → 符合或优于
                if d_val > r_val:
                    return "优于"
                elif d_val == r_val:
                    return "符合"
                else:
                    return "不符合"
            elif d_dir == 'max':
                # 最大值约束：草稿的最大值 <= 引用标准的最大值 → 符合或优于
                if d_val < r_val:
                    return "优于"
                elif d_val == r_val:
                    return "符合"
                else:
                    return "不符合"
            else:
                # 等于约束
                if d_val == r_val:
                    return "符合"
                else:
                    return "无法判断"
        else:
            # 不同方向约束，无法直接比较
            return "无法判断"
    
    # ============================================================
    # 辅助方法
    # ============================================================
    
    def _safe_add_issue(self, code: str, severity: str, location: str,
                        description: str, suggestion: str, context: str = "") -> None:
        """安全添加问题（如果有回调函数）"""
        if self._add_issue:
            # 将字符串 severity 转换为 Severity 枚举
            from enum import Enum
            
            class Sev(Enum):
                ERROR = "ERROR"
                WARNING = "WARNING"
                SUGGESTION = "SUGGESTION"
            
            sev = Sev[severity] if isinstance(severity, str) else severity
            self._add_issue(code, sev, location, description, suggestion, context)
    
    def get_summary(self) -> Dict[str, Any]:
        """获取检查摘要"""
        return {
            "total_references": len(self.references),
            "total_citations": len(self.citations),
            "draft_requirements": len(self.draft_requirements),
            "ref_requirements": len(self.ref_requirements),
            "matched_pairs": self.matched_pairs_count,
            "references_detail": [
                {
                    "number": r.number,
                    "title": r.title,
                    "is_dated": r.is_dated,
                    "location": r.location,
                }
                for r in self.references
            ],
        }
