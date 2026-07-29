#!/usr/bin/env python3
"""测试脚本：创建带已知问题的样例 .docx，验证新增规则检测"""

import sys
import os
import json

# 添加脚本路径
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from lxml import etree


def create_test_docx(path):
    """创建包含各种已知问题的测试文档"""
    doc = Document()

    # === 目次 (F010: 页码用阿拉伯数字) ===
    doc.add_heading("目次", level=1)
    p = doc.add_paragraph("前言..............1")
    p = doc.add_paragraph("范围..............3")

    # === 前言 ===
    doc.add_heading("前言", level=1)
    doc.add_paragraph("本文件按照GB/T 1.1-2020的规定起草。")

    # === 引言 (S009: 引言包含"应") ===
    doc.add_heading("引言", level=1)
    doc.add_paragraph("本文件的制定应满足行业发展的需求。")  # S009: 引言中有"应"

    # === 范围 ===
    doc.add_heading("1 范围", level=1)
    doc.add_paragraph("本文件规定了某产品的技术要求。")
    doc.add_paragraph("本文件适用于某产品的生产和使用。")

    # === 规范性引用文件 (S007: 无内容未声明) ===
    doc.add_heading("2 规范性引用文件", level=1)
    # 故意留空，不写声明语句

    # === 术语和定义 ===
    doc.add_heading("3 术语和定义", level=1)

    # W003: 术语未加粗
    p = doc.add_paragraph()
    run = p.add_run("可靠性")
    run.bold = False  # 故意不加粗
    p.add_run("是指产品在规定条件下和规定时间内完成规定功能的能力。")

    # T009: 术语定义含要求型条款
    p = doc.add_paragraph()
    run = p.add_run("安全设备")
    run.bold = True
    p.add_run("是指在运行中应始终保持安全状态的设备。")  # T009: 术语定义中有"应"

    # === 技术要求 ===
    doc.add_heading("4 技术要求", level=1)

    # S004: 孤立条编号（只有4.1无4.2）
    doc.add_heading("4.1 外观要求", level=2)
    doc.add_paragraph("产品外观应平整光滑。")

    # F003: 尺寸表述不规范
    doc.add_paragraph("产品外形尺寸为 80x25x50 mm。")  # F003

    # F011: 公式编号无括号
    doc.add_paragraph("强度计算见公式3。")  # F011

    # W001: 中文正文使用半角标点
    doc.add_paragraph("产品应符合下列要求:重量不超过50kg,长度不小于1m;")  # W001 (半角:和,和;)

    # T004: 遵守/符合混用
    doc.add_paragraph("产品应遵守以下技术要求。")  # T004: 遵守...要求

    # T008: 概述含要求条款
    doc.add_heading("4.2 概述", level=2)
    doc.add_paragraph("本节应规定产品的通用技术要求。")  # T008: 概述中有"应"

    # S005: 破折号引导的款直接出现在章下
    doc.add_heading("5 测试方法", level=1)
    p = doc.add_paragraph("— 测试应在标准环境下进行。")  # S005: 破折号直接在章下

    # S004: 孤立条编号（只有6.1无6.2）
    doc.add_heading("6 检验规则", level=1)
    doc.add_heading("6.1 出厂检验", level=2)
    doc.add_paragraph("出厂检验应逐台进行。")
    # 注意：这里故意不添加 6.2

    # S006: 附录编号使用I
    doc.add_heading("附录I 补充数据", level=1)  # S006

    # F012: 图/表脚注使用数字编号
    doc.add_paragraph("图1 注1: 此处为图脚注说明。")  # F012

    doc.save(path)
    print(f"测试文档已创建: {path}")


def run_tests():
    """运行测试"""
    from check_standard import StandardChecker

    test_docx = os.path.join(os.path.dirname(os.path.abspath(__file__)), "test_sample.docx")
    create_test_docx(test_docx)

    checker = StandardChecker()
    issues = checker.check(test_docx, analyze_styles=False)

    # 预期检测到的规则
    expected_codes = {
        "F003": "尺寸表述不规范",
        "F010": "目次使用阿拉伯数字",
        "F011": "公式编号无括号",
        "F012": "图表脚注编号格式错误",
        "S004": "孤立条编号",
        "S005": "款/项无归属",
        "S006": "附录编号使用I或O",
        "S007": "无内容要素未声明",
        "S009": "引言包含规范性条款",
        "T004": "遵守/符合混用",
        "T008": "概述含要求条款",
        "T009": "术语定义含要求型条款",
        "W001": "中文正文使用半角标点",
        "W003": "术语首次出现未加粗",
    }

    # 收集实际检测到的规则
    found_codes = set()
    for issue in issues:
        found_codes.add(issue.code)

    # 汇报结果
    print("\n" + "=" * 60)
    print("测试结果")
    print("=" * 60)

    passed = 0
    failed = 0

    for code, desc in expected_codes.items():
        if code in found_codes:
            print(f"  ✓ {code} - {desc}")
            passed += 1
        else:
            print(f"  ✗ {code} - {desc} [未检测到]")
            failed += 1

    # 显示额外检测到的问题（不在预期中的）
    extra_codes = found_codes - set(expected_codes.keys())
    # 排除原有规则
    original_codes = {"S001", "S002", "S003", "F001", "F002", "F004", "F005", "F006",
                      "F007", "F008", "F009", "T001", "T002", "T003", "T005", "T006"}
    extra_new = extra_codes - original_codes
    if extra_new:
        print(f"\n额外检测到的新规则:")
        for code in extra_new:
            issue = next(i for i in issues if i.code == code)
            print(f"  + {code} - {issue.description}")

    print(f"\n总计: {passed} 通过, {failed} 失败")
    print(f"总检测问题数: {len(issues)}")

    # 打印所有问题的详情
    print("\n" + "-" * 60)
    print("所有检测到的问题:")
    print("-" * 60)
    for issue in issues:
        print(f"  [{issue.severity}] {issue.code}: {issue.description} ({issue.location})")

    # 清理测试文件
    try:
        os.remove(test_docx)
        print(f"\n测试文件已清理。")
    except Exception:
        print(f"\n测试文件清理失败（沙箱限制），文件位置: {test_docx}")

    return failed == 0


if __name__ == "__main__":
    success = run_tests()
    sys.exit(0 if success else 1)
