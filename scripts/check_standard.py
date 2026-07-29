#!/usr/bin/env python3
"""
GB/T 1.1-2020 标准草稿自动化检查脚本（改进版）

功能：
1. 分析文档样式体系，智能识别标题
2. 执行结构化规则检查（编号连续性、能愿动词、格式等）
3. 输出 JSON 格式的检查结果

使用方法：
    python check_standard.py <input.docx|input.pdf> [--output result.json] [--analyze-styles]
"""

import re
import sys
import json
import argparse
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, asdict
from enum import Enum
from collections import Counter

# 多标准支持
from standard_profiles import (
    StandardProfile,
    get_profile,
    list_profiles,
    auto_detect,
    is_rule_enabled,
)

# 引用标准符合性检查
from reference_checker import ReferenceChecker


class Severity(Enum):
    ERROR = "ERROR"
    WARNING = "WARNING"
    SUGGESTION = "SUGGESTION"


@dataclass
class Issue:
    """检查问题"""
    code: str          # 问题编号，如 S001, F001
    severity: str      # 严重等级：ERROR/WARNING/SUGGESTION
    location: str      # 位置描述
    description: str   # 问题描述
    suggestion: str    # 修改建议
    context: str = ""  # 上下文信息（可选）


class StyleAnalyzer:
    """文档样式分析器"""
    
    def __init__(self):
        self.style_stats: Dict[str, Dict] = {}
        self.heading_styles: List[str] = []
        
    def analyze(self, doc) -> Dict[str, Any]:
        """分析文档样式"""
        style_counter = Counter()
        style_samples: Dict[str, List[str]] = {}
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            style_name = para.style.name if para.style else "None"
            style_counter[style_name] += 1
            
            # 收集样式样本
            if style_name not in style_samples:
                style_samples[style_name] = []
            if len(style_samples[style_name]) < 3:
                style_samples[style_name].append(text[:80])
        
        # 分析样式特征
        for style_name, count in style_counter.items():
            samples = style_samples.get(style_name, [])
            
            # 判断是否为标题样式
            is_heading = self._is_heading_style(style_name, samples)
            
            self.style_stats[style_name] = {
                "count": count,
                "samples": samples,
                "is_heading": is_heading
            }
            
            if is_heading:
                self.heading_styles.append(style_name)
        
        return {
            "style_stats": self.style_stats,
            "heading_styles": self.heading_styles,
            "total_styles": len(style_counter)
        }
    
    def _is_heading_style(self, style_name: str, samples: List[str]) -> bool:
        """判断样式是否为标题样式"""
        
        # 1. 明确的标题样式关键词
        heading_keywords = ["章标题", "一级条标题", "标题 1", "Heading 1", "标题"]
        
        # 2. 明确的非标题样式关键词
        non_heading_keywords = ["无标题条", "正文", "Body Text", "Normal"]
        
        # 检查样式名称
        for keyword in heading_keywords:
            if keyword in style_name:
                # 但要排除"无标题条"
                if "无标题条" not in style_name:
                    return True
        
        for keyword in non_heading_keywords:
            if keyword in style_name:
                return False
        
        # 3. 根据样本内容判断
        if samples:
            # 如果样本都是长文本（>50字符）且以句号结尾，可能是正文
            avg_length = sum(len(s) for s in samples) / len(samples)
            ends_with_period = all(s.endswith('。') for s in samples if s)
            
            if avg_length > 50 and ends_with_period:
                return False
            
            # 如果样本符合编号格式，可能是标题
            for sample in samples:
                if re.match(r'^\d+\s+\S', sample) or re.match(r'^\d+\.\d+\s+\S', sample):
                    if len(sample) <= 50:
                        return True
        
        return False


class StandardChecker:
    """标准草稿检查器（支持多标准类型）"""

    def __init__(self, standard: str = None):
        """
        Args:
            standard: 标准类型 ID。None 时在 check() 中自动检测。
        """
        self.issues: List[Issue] = []
        self.paragraphs: List[Dict] = []
        self.headings: List[Dict] = []
        self.style_analyzer = StyleAnalyzer()
        self.doc = None
        self.section_map: Dict[int, str] = {}
        # 多标准配置
        self.profile: Optional[StandardProfile] = None
        self._pending_standard = standard  # 延迟到段落提取后再解析
        
    def check(self, file_path: str, analyze_styles: bool = False,
              ref_files: List[str] = None) -> List[Issue]:
        """执行检查（支持 .docx 和 .pdf 格式）

        Args:
            file_path: 输入文件路径
            analyze_styles: 是否分析文档样式
            ref_files: 引用标准文档路径列表（用于符合性检查，可选）
        """
        input_path = Path(file_path)
        if not input_path.exists():
            print(f"错误：文件不存在：{input_path}")
            sys.exit(1)

        file_ext = input_path.suffix.lower()

        if file_ext == '.pdf':
            # PDF 文件：使用 pdf_extractor
            try:
                from pdf_extractor import extract as extract_pdf
            except ImportError:
                print("错误：需要安装 PyMuPDF 库")
                print("请运行：pip install PyMuPDF")
                sys.exit(1)

            doc = extract_pdf(str(input_path))
            self.doc = doc
            print(f"已加载 PDF 文件：{input_path}（{len(doc.paragraphs)} 个段落）")
        else:
            # DOCX 文件
            try:
                from docx import Document
            except ImportError:
                print("错误：需要安装 python-docx 库")
                print("请运行：pip install python-docx")
                sys.exit(1)

            doc = Document(str(input_path))
            self.doc = doc

            # 分析样式（可选）
            if analyze_styles:
                style_info = self.style_analyzer.analyze(doc)
                print("\n=== 文档样式分析 ===")
                print(f"总样式数: {style_info['total_styles']}")
                print(f"标题样式: {', '.join(style_info['heading_styles'])}")
        
        # 提取文档结构
        self._extract_structure(doc)

        # 确定标准类型 profile
        if self._pending_standard:
            self.profile = get_profile(self._pending_standard)
        else:
            detected = auto_detect(self.paragraphs)
            self.profile = get_profile(detected)
        print(f"标准类型：{self.profile.name}（{self.profile.id}）")
        print(f"起草依据：{self.profile.drafting_standard}")

        # 执行各项检查
        self._check_structure()
        self._check_headings()
        self._check_modal_verbs()
        self._check_number_unit_spacing()
        self._check_figure_table_numbering()
        self._check_references()
        self._check_percentage_tolerance()
        self._check_hyphen_consistency()  # 连字符一致性检查

        # 新增规则检查
        self._check_dimension_format()             # F003
        self._check_formula_numbering()            # F011
        self._check_isolated_clauses()             # S004
        self._check_appendix_numbering()           # S006
        self._check_introduction_normative()       # S009
        self._check_summary_normative()            # T008
        self._check_chinese_punctuation()          # W001
        self._check_toc_numbering()                # F010
        self._check_figure_table_footnotes()       # F012
        self._check_dashed_paragraphs()            # S005
        self._check_empty_sections()               # S007
        self._check_comply_conform()               # T004
        self._check_footnote_requirements()        # T007
        self._check_term_definition_requirements() # T009
        self._check_term_bold()                    # W003

        # 标准类型专属检查
        for check_name in self.profile.specific_checks:
            method = getattr(self, check_name, None)
            if method:
                method()

        # 引用标准符合性检查
        self.ref_checker = ReferenceChecker(
            self.paragraphs, self.headings, self.section_map,
            add_issue_func=self._add_issue
        )
        self.ref_checker.check_auto()

        # 用户提交层：符合性检查（仅当提供了引用标准文档时）
        if ref_files:
            self.ref_checker.check_compliance(ref_files)

        return self.issues
    
    def _extract_structure(self, doc) -> None:
        """提取文档结构"""
        current_section = ""
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""

            # 智能识别标题
            is_heading = self._is_heading(text, style_name)

            # 更新当前章节
            if is_heading:
                current_section = self._get_section_name(text)

            item = {
                "index": i,
                "text": text,
                "style": style_name,
                "is_heading": is_heading,
                "section": current_section
            }

            self.paragraphs.append(item)
            self.section_map[i] = current_section

            if is_heading:
                self.headings.append(item)
    
    def _is_heading(self, text: str, style_name: str) -> bool:
        """智能判断是否为标题"""

        # 1. 使用样式分析器的结果
        if style_name in self.style_analyzer.heading_styles:
            # 额外检查：长度和内容
            if len(text) <= 50 and not text.endswith('。'):
                return True

        # 1b. 检查样式名称是否包含 Heading 或 标题（无需样式分析也能识别）
        if style_name and ("Heading" in style_name or "标题" in style_name):
            if "无标题" not in style_name and len(text) <= 50 and not text.endswith('。'):
                return True

        # 2. 检查编号格式
        # 章编号：数字+空格+文字（如"1 范围"）
        if re.match(r'^\d+\s+\S', text):
            # 排除目录项（包含制表符）
            if '\t' not in text:
                # 长度检查
                if len(text) <= 50:
                    return True

        # 条编号：数字.数字+空格+文字（如"5.1 总体要求"）
        if re.match(r'^\d+\.\d+\s+\S', text):
            if '\t' not in text and len(text) <= 50:
                return True

        # 3. 检查已知非编号标题（前言、引言、目次等）
        # 排除目次条目（含连续点号或省略号+页码）
        if not re.search(r'\.{2,}|…', text):
            known_headings = ["前言", "引言", "目次", "参考文献", "索引", "概述"]
            for kh in known_headings:
                if text == kh or text.startswith(kh):
                    if len(text) <= 20 and not text.endswith('。'):
                        return True

        # 4. 检查附录标题
        if re.match(r'^附录\s*[A-Z]', text) and len(text) <= 50:
            return True

        return False
    
    def _add_issue(self, code: str, severity: Severity, location: str,
                   description: str, suggestion: str, context: str = "") -> None:
        """添加问题（自动按 profile 过滤禁用规则）"""
        # 按 profile 过滤
        if self.profile and not is_rule_enabled(self.profile, code):
            return
        self.issues.append(Issue(
            code=code,
            severity=severity.value,
            location=location,
            description=description,
            suggestion=suggestion,
            context=context
        ))
    
    def _check_structure(self) -> None:
        """检查结构要素"""
        # 从 profile 获取必备要素
        required_elements = self.profile.required_elements if self.profile else ["前言", "范围"]
        heading_texts = [h["text"] for h in self.headings]
        all_text = " ".join([p["text"] for p in self.paragraphs])
        
        for elem in required_elements:
            # 改进的识别逻辑：移除空格后检查
            found_in_heading = False
            found_in_text = False
            
            # 检查标题（考虑空格）
            for h in heading_texts:
                # 移除所有类型的空格（半角、全角、特殊空格）
                h_no_space = h.replace(' ', '').replace('　', '').replace('\u2002', '').replace('\u2003', '')
                if elem in h_no_space or h_no_space == elem:
                    found_in_heading = True
                    # 检查标题是否包含空格
                    if ' ' in h or '　' in h or '\u2002' in h or '\u2003' in h:
                        self._add_issue(
                            "F008",
                            Severity.WARNING,
                            f"标题",
                            f"\"{elem}\"标题包含装饰性空格",
                            "标题文字内部不应包含空格",
                            h
                        )
                    break
            
            # 检查全文
            for p in self.paragraphs:
                text_no_space = p["text"].replace(' ', '').replace('　', '')
                if elem in text_no_space:
                    found_in_text = True
                    break
            
            if not found_in_text:
                self._add_issue(
                    "S001",
                    Severity.ERROR,
                    "整体",
                    f"缺少必备要素\"{elem}\"",
                    f"补充\"{elem}\"要素"
                )
        
        # 检查要素顺序
        element_order = []
        for h in self.headings:
            text = h["text"]
            if "前言" in text:
                element_order.append("前言")
            elif "范围" in text:
                element_order.append("范围")
            elif "规范性引用文件" in text:
                element_order.append("规范性引用文件")
        
        expected_order = ["前言", "范围", "规范性引用文件"]
        for i, elem in enumerate(expected_order):
            if elem in element_order:
                idx = element_order.index(elem)
                for j in range(i):
                    if expected_order[j] in element_order:
                        if element_order.index(expected_order[j]) > idx:
                            self._add_issue(
                                "S002",
                                Severity.ERROR,
                                "整体",
                                f"要素顺序错误：\"{elem}\"应在\"{expected_order[j]}\"之后",
                                "调整要素顺序为：前言->范围->规范性引用文件->..."
                            )
                            break
    
    def _check_headings(self) -> None:
        """检查标题格式"""
        chapter_pattern = re.compile(r'^(\d+)\s+')
        clause_pattern = re.compile(r'^(\d+\.\d+)\s+')
        
        chapter_numbers = []
        clause_numbers = []
        
        for h in self.headings:
            text = h["text"]
            
            # 检查标题末尾标点
            if text and text[-1] in '：:，,。.、；;!！':
                self._add_issue(
                    "F001",
                    Severity.ERROR,
                    f"位置 {h['index']}",
                    f"标题末尾有标点：\"{text}\"",
                    "删除末尾标点",
                    f"样式: {h['style']}"
                )
            
            # 提取章编号
            chapter_match = chapter_pattern.match(text)
            if chapter_match:
                num = int(chapter_match.group(1))
                chapter_numbers.append(num)
            
            # 提取条编号
            clause_match = clause_pattern.match(text)
            if clause_match:
                clause_numbers.append(clause_match.group(1))
        
        # 检查章编号连续性
        if chapter_numbers:
            chapter_numbers = sorted(set(chapter_numbers))
            for i in range(len(chapter_numbers) - 1):
                if chapter_numbers[i+1] - chapter_numbers[i] > 1:
                    self._add_issue(
                        "S003",
                        Severity.ERROR,
                        "整体",
                        f"章编号不连续：从{chapter_numbers[i]}跳到{chapter_numbers[i+1]}",
                        "补充缺失章节或重新编号"
                    )
    
    def _check_modal_verbs(self) -> None:
        """检查能愿动词使用"""
        # 禁用词
        forbidden_words = ["必须", "应当", "不得"]
        # 弱化词
        weakening_words = ["尽量", "尽可能", "考虑", "优先考虑", "充分考虑", "避免", "慎重"]
        # 限定词
        limiting_words = ["通常", "一般", "原则上"]
        
        for para in self.paragraphs:
            text = para["text"]
            
            # 检查禁用词
            for word in forbidden_words:
                if word in text:
                    replacement = "应" if word != "不得" else "不应"
                    self._add_issue(
                        "T001",
                        Severity.ERROR,
                        f"段落 {para['index']}",
                        f"使用禁用能愿动词\"{word}\"",
                        f"改为\"{replacement}\"",
                        text[:100]
                    )
            
            # 检查"应"与弱化词搭配
            if "应" in text:
                for word in weakening_words:
                    if word in text:
                        self._add_issue(
                            "T002",
                            Severity.ERROR,
                            f"段落 {para['index']}",
                            f"\"应\"与弱化词\"{word}\"搭配",
                            f"改为\"宜{word}\"",
                            text[:100]
                        )
            
            # 检查"应"与限定词搭配
            for word in limiting_words:
                if f"{word}应" in text or f"{word} 应" in text:
                    self._add_issue(
                        "T003",
                        Severity.WARNING,
                        f"段落 {para['index']}",
                        f"\"应\"与限定词\"{word}\"搭配",
                        f"改为\"{word}宜\"",
                        text[:100]
                    )
    
    def _check_number_unit_spacing(self) -> None:
        """检查数字与单位间距"""
        # 常见单位
        units = ["mm", "cm", "m", "km", "g", "kg", "t", "ml", "L", "W", "kW",
                 "V", "A", "Hz", "kHz", "MHz", "Pa", "kPa", "MPa", "J", "kJ",
                 "C", "K", "min", "h", "s", "ms", "r", "r/min"]
        
        for para in self.paragraphs:
            text = para["text"]
            
            # 检查数字+单位无空格
            for unit in units:
                # 匹配数字直接接单位（无空格）
                pattern = re.compile(rf'(\d){unit}(?![a-zA-Z])')
                if pattern.search(text):
                    self._add_issue(
                        "F002",
                        Severity.WARNING,
                        f"段落 {para['index']}",
                        f"数字与单位\"{unit}\"之间缺少空格",
                        "在数字与单位间添加空格"
                    )
    
    def _check_figure_table_numbering(self) -> None:
        """检查图表编号"""
        # 提取图编号
        figure_pattern = re.compile(r'图\s*(\d+)')
        table_pattern = re.compile(r'表\s*(\d+)')
        
        figure_numbers = []
        table_numbers = []
        
        for para in self.paragraphs:
            text = para["text"]
            
            # 提取图编号
            for match in figure_pattern.finditer(text):
                figure_numbers.append(int(match.group(1)))
            
            # 提取表编号
            for match in table_pattern.finditer(text):
                table_numbers.append(int(match.group(1)))
        
        # 检查图编号连续性
        if figure_numbers:
            figure_numbers = sorted(set(figure_numbers))
            for i in range(len(figure_numbers) - 1):
                if figure_numbers[i+1] - figure_numbers[i] > 1:
                    self._add_issue(
                        "F005",
                        Severity.ERROR,
                        "整体",
                        f"图编号不连续：从图{figure_numbers[i]}跳到图{figure_numbers[i+1]}",
                        "重新编号或补充缺失图"
                    )
        
        # 检查表编号连续性
        if table_numbers:
            table_numbers = sorted(set(table_numbers))
            for i in range(len(table_numbers) - 1):
                if table_numbers[i+1] - table_numbers[i] > 1:
                    self._add_issue(
                        "F006",
                        Severity.ERROR,
                        "整体",
                        f"表编号不连续：从表{table_numbers[i]}跳到表{table_numbers[i+1]}",
                        "重新编号或补充缺失表"
                    )
        
        # 检查分表编号
        for para in self.paragraphs:
            text = para["text"]
            if re.search(r'表\d+[a-zA-Z]', text):
                self._add_issue(
                    "F007",
                    Severity.ERROR,
                    f"段落 {para['index']}",
                    "存在分表编号（如表2a）",
                    "合并为一个表或分别编号"
                )
    
    def _check_references(self) -> None:
        """检查引用格式"""
        # 注日期引用格式（应为 GB/T XXXXX-YYYY，用一字线）
        dated_ref_pattern = re.compile(r'GB/T\s*\d+\.?\d*[:：]\d{4}')
        # 页码引用
        page_ref_pattern = re.compile(r'第\s*\d+\s*页')
        
        for para in self.paragraphs:
            text = para["text"]
            
            # 检查注日期引用格式
            if dated_ref_pattern.search(text):
                self._add_issue(
                    "T006",
                    Severity.ERROR,
                    f"段落 {para['index']}",
                    "注日期引用格式错误（使用了冒号）",
                    "使用一字线连接，如 GB/T XXXXX-2020"
                )
            
            # 检查页码引用
            if page_ref_pattern.search(text):
                self._add_issue(
                    "T005",
                    Severity.ERROR,
                    f"段落 {para['index']}",
                    "引用中包含页码",
                    "引用内容编号而非页码，如\"第3章\"\"5.2条\""
                )
    
    def _check_percentage_tolerance(self) -> None:
        """检查百分率公差格式"""
        # 匹配错误的百分率公差格式
        wrong_tolerance_pattern = re.compile(r'\d+\s*[+＋]\s*[-－]\s*\d+\s*%')
        
        for para in self.paragraphs:
            text = para["text"]
            
            if wrong_tolerance_pattern.search(text):
                self._add_issue(
                    "F004",
                    Severity.ERROR,
                    f"段落 {para['index']}",
                    "百分率公差格式错误",
                    "使用括号包裹，如 (65+/-2)%"
                )
    
    def _check_hyphen_consistency(self) -> None:
        """检查连字符一致性（新增）"""
        # 检查标准引用中的连字符格式
        # 半角连字符模式：GB/T 20258.2-2019
        half_width_pattern = re.compile(r'GB/T\s+\d+\.?\d*-(\d{4})')
        # 全角连字符模式：GB/T 1.1—2020
        full_width_pattern = re.compile(r'GB/T\s+\d+\.?\d*—(\d{4})')
        
        half_width_refs = []
        full_width_refs = []
        
        for para in self.paragraphs:
            text = para["text"]
            
            # 查找半角连字符
            for match in half_width_pattern.finditer(text):
                half_width_refs.append({
                    "text": match.group(),
                    "location": f"段落 {para['index']}"
                })
            
            # 查找全角连字符
            for match in full_width_pattern.finditer(text):
                full_width_refs.append({
                    "text": match.group(),
                    "location": f"段落 {para['index']}"
                })
        
        # 如果同时存在半角和全角连字符，报告不一致
        if half_width_refs and full_width_refs:
            # 构建上下文信息
            half_examples = [ref["text"] for ref in half_width_refs[:3]]
            full_examples = [ref["text"] for ref in full_width_refs[:3]]
            
            context = f"半角示例：{', '.join(half_examples)}\n全角示例：{', '.join(full_examples)}"
            
            self._add_issue(
                "F009",
                Severity.WARNING,
                "整体",
                f'标准引用中的连字符格式不一致：{len(half_width_refs)}处使用半角"-"，{len(full_width_refs)}处使用全角"—"',
                '统一使用半角连字符"-"（推荐）或全角连字符"—"',
                context
            )

    # ===== 标准类型专属检查 =====

    def _check_national_standard_number(self) -> None:
        """MS001: 国家标准编号格式检查"""
        all_text = " ".join(p["text"] for p in self.paragraphs)
        # GB/T XXXXX-YYYY 或 GB XXXXX-YYYY
        valid_pattern = re.compile(r'GB/T?\s*\d+\.?\d*[-—]\d{4}')
        # 错误格式：使用冒号
        wrong_pattern = re.compile(r'GB/T?\s*\d+\.?\d*[:：]\d{4}')

        for para in self.paragraphs:
            text = para["text"]
            if wrong_pattern.search(text):
                self._add_issue(
                    "MS001",
                    Severity.ERROR,
                    f"段落 {para['index']}",
                    "国家标准编号使用冒号连接年份",
                    "使用一字线，如 GB/T 12345-2020",
                    text[:100]
                )

    def _check_industry_standard_number(self) -> None:
        """MS002: 行业标准编号格式检查"""
        # 常见行业标准前缀
        known_prefixes = {
            "YY": "医药", "JB": "机械", "QB": "轻工", "CJ": "城镇建设",
            "SL": "水利", "JG": "建筑工业", "GA": "公共安全", "HG": "化工",
            "SY": "石油天然气", "DL": "电力", "TB": "铁路", "JT": "交通",
            "TD": "土地管理", "LY": "林业", "NY": "农业", "SB": "商务",
            "WS": "卫生", "YZ": "邮政", "AQ": "安全生产", "MT": "煤炭",
            "YS": "有色金属", "EJ": "核工业", "QX": "气象", "DZ": "地震",
            "GH": "城乡建设规划", "HJ": "环境保护", "DB": "地震（非地方标准）",
            "CB": "船舶", "CH": "测绘", "CY": "新闻出版", "DA": "档案",
            "FZ": "纺织", "GY": "广播影视", "HB": "航空", "JC": "建材",
            "JJ": "计量", "SH": "石油化工", "SJ": "电子", "SN": "商检",
            "WH": "文化", "WJ": "兵工民品", "WM": "外经贸", "XB": "稀土",
            "ZB": "专业标准",
        }

        for para in self.paragraphs:
            text = para["text"]
            # 匹配行业标准编号
            m = re.match(r'^([A-Z]{2})/T?\s*\d+', text)
            if m:
                prefix = m.group(1)
                if prefix == "GB":
                    continue
                if prefix not in known_prefixes:
                    self._add_issue(
                        "MS002",
                        Severity.WARNING,
                        f"段落 {para['index']}",
                        f"行业标准编号前缀\"{prefix}\"不在已知行业代号列表中",
                        "核实行业代号是否正确",
                        text[:100]
                    )

    def _check_local_standard_number(self) -> None:
        """MS003: 地方标准编号格式检查"""
        # DBXX/T XXXXX-YYYY，XX 为行政区划代码
        # 有效的省级行政区划代码
        valid_codes = {
            "11": "北京", "12": "天津", "13": "河北", "14": "山西", "15": "内蒙古",
            "21": "辽宁", "22": "吉林", "23": "黑龙江",
            "31": "上海", "32": "江苏", "33": "浙江", "34": "安徽", "35": "福建", "36": "江西",
            "37": "山东",
            "41": "河南", "42": "湖北", "43": "湖南", "44": "广东", "45": "广西", "46": "海南",
            "50": "重庆", "51": "四川", "52": "贵州", "53": "云南", "54": "西藏",
            "61": "陕西", "62": "甘肃", "63": "青海", "64": "宁夏", "65": "新疆",
            "71": "中国台湾", "81": "中国香港", "82": "中国澳门",
        }

        for para in self.paragraphs:
            text = para["text"]
            # 匹配 DB + 2位数字 + /T?
            m = re.search(r'DB(\d{2})/?T?\s*\d+', text)
            if m:
                code = m.group(1)
                if code not in valid_codes:
                    self._add_issue(
                        "MS003",
                        Severity.ERROR,
                        f"段落 {para['index']}",
                        f"地方标准区域代码\"{code}\"无效",
                        "使用正确的省级行政区划代码（如 11=北京, 31=上海, 44=广东）",
                        text[:100]
                    )

    def _check_enterprise_standard_number(self) -> None:
        """MS004: 企业标准编号格式检查"""
        for para in self.paragraphs:
            text = para["text"]
            # Q/XXXXXXXX-YYYY
            # 检查是否缺少年份
            m = re.search(r'Q/([A-Z]+)\s*\d+', text)
            if m:
                # 检查是否有年份
                if not re.search(r'Q/[A-Z]+\s*\d+[-—:：]\d{4}', text):
                    self._add_issue(
                        "MS004",
                        Severity.WARNING,
                        f"段落 {para['index']}",
                        "企业标准编号缺少年份",
                        "补充年份，如 Q/ABC 001-2023",
                        text[:100]
                    )

    def _check_group_standard_number(self) -> None:
        """MS005: 团体标准编号格式检查"""
        for para in self.paragraphs:
            text = para["text"]
            # T/XXXXXXXX-YYYY
            # 检查是否缺少年份
            m = re.search(r'T/([A-Z]+)\s*\d+', text)
            if m:
                # 检查是否有年份
                if not re.search(r'T/[A-Z]+\s*\d+[-—:：]\d{4}', text):
                    self._add_issue(
                        "MS005",
                        Severity.WARNING,
                        f"段落 {para['index']}",
                        "团体标准编号缺少年份",
                        "补充年份，如 T/CAS 001-2023",
                        text[:100]
                    )

    # ===== 辅助方法 =====

    def _get_section_name(self, heading_text: str) -> str:
        """从标题文本中提取章节名称"""
        text = re.sub(r'^\d+(\.\d+)*\s+', '', heading_text)
        text = text.strip()
        if "前言" in text:
            return "前言"
        elif "引言" in text:
            return "引言"
        elif "范围" in text:
            return "范围"
        elif "规范性引用文件" in text or "引用文件" in text:
            return "规范性引用文件"
        elif "术语" in text and "定义" in text:
            return "术语和定义"
        elif "概述" in text:
            return "概述"
        elif "目次" in text:
            return "目次"
        elif re.match(r'^附录', text):
            return "附录"
        else:
            return text

    # ===== 简单难度规则 =====

    def _check_dimension_format(self) -> None:
        """F003: 检查尺寸表述不规范（如 80x25x50 mm → 80 mm x 25 mm x 50 mm）"""
        units = r'(?:mm|cm|m|km|μm|um|nm)'
        # 匹配 数字x数字(+x数字) 后面跟单位，且数字间无单位
        pattern = re.compile(rf'(\d+(?:\.\d+)?)\s*[x×]\s*(\d+(?:\.\d+)?)(?:\s*[x×]\s*(\d+(?:\.\d+)?))?\s*{units}')

        for para in self.paragraphs:
            text = para["text"]
            for match in pattern.finditer(text):
                full = match.group()
                # 如果第一个数字和x之间有单位，说明是规范的（如 80 mm x 25 mm）
                if re.match(rf'\d+(?:\.\d+)?\s*{units}\s*[x×]', full):
                    continue
                self._add_issue(
                    "F003",
                    Severity.WARNING,
                    f"段落 {para['index']}",
                    f"尺寸表述不规范：\"{full}\"",
                    "每个量后须带单位，如 80 mm x 25 mm x 50 mm",
                    text[:100]
                )

    def _check_formula_numbering(self) -> None:
        """F011: 检查公式编号是否有括号（如 公式3 → 公式(3)）"""
        # 匹配"公式"后直接跟数字（无括号）
        pattern = re.compile(r'公式\s*(?!\()(\d+)')

        for para in self.paragraphs:
            text = para["text"]
            for match in pattern.finditer(text):
                self._add_issue(
                    "F011",
                    Severity.WARNING,
                    f"段落 {para['index']}",
                    f"公式编号无括号：\"公式{match.group(1)}\"",
                    f"改为\"公式({match.group(1)})\"",
                    text[:100]
                )

    def _check_isolated_clauses(self) -> None:
        """S004: 检查孤立条编号（只有3.1无3.2）"""
        clause_map: Dict[int, List[int]] = {}

        clause_pattern = re.compile(r'^(\d+)\.(\d+)\s+')
        for h in self.headings:
            match = clause_pattern.match(h["text"])
            if match:
                chapter = int(match.group(1))
                clause = int(match.group(2))
                if chapter not in clause_map:
                    clause_map[chapter] = []
                if clause not in clause_map[chapter]:
                    clause_map[chapter].append(clause)

        for chapter, clauses in clause_map.items():
            if len(clauses) == 1:
                self._add_issue(
                    "S004",
                    Severity.WARNING,
                    "整体",
                    f"孤立条编号：只有{chapter}.1而无{chapter}.2",
                    "改为在章下设段，或补充其他条",
                    f"章{chapter}仅有条{chapter}.1"
                )

    def _check_appendix_numbering(self) -> None:
        """S006: 检查附录编号是否使用了I或O"""
        pattern = re.compile(r'附录\s*([A-Z])')

        for para in self.paragraphs:
            text = para["text"]
            for match in pattern.finditer(text):
                letter = match.group(1)
                if letter in ('I', 'O'):
                    self._add_issue(
                        "S006",
                        Severity.ERROR,
                        f"段落 {para['index']}",
                        f"附录编号使用了\"{letter}\"（易与数字混淆）",
                        f"跳过字母{letter}，使用前一个或后一个字母",
                        text[:100]
                    )

    def _check_introduction_normative(self) -> None:
        """S009: 检查引言是否包含规范性条款（"应"）"""
        intro_paras = [p for p in self.paragraphs
                       if p.get("section") == "引言" and not p["is_heading"]]

        for para in intro_paras:
            text = para["text"]
            if re.search(r'应[\u4e00-\u9fff]', text):
                self._add_issue(
                    "S009",
                    Severity.ERROR,
                    f"段落 {para['index']}",
                    "引言包含规范性条款（出现\"应\"）",
                    "将规范性内容移至正文，引言只能含资料性内容",
                    text[:100]
                )

    def _check_summary_normative(self) -> None:
        """T008: 检查"概述"章节是否含要求条款"""
        summary_paras = [p for p in self.paragraphs
                         if p.get("section") == "概述" and not p["is_heading"]]

        for para in summary_paras:
            text = para["text"]
            if re.search(r'应[\u4e00-\u9fff]', text):
                self._add_issue(
                    "T008",
                    Severity.ERROR,
                    f"段落 {para['index']}",
                    "\"概述\"章节含要求条款（出现\"应\"）",
                    "\"概述\"只能含陈述型条款",
                    text[:100]
                )

    def _check_chinese_punctuation(self) -> None:
        """W001: 检查中文正文是否使用半角标点"""
        half_to_full = {
            ',': '，',
            ':': '：',
            ';': '；',
        }

        for para in self.paragraphs:
            text = para["text"]
            if not text or para["is_heading"]:
                continue

            for i, char in enumerate(text):
                if char in half_to_full:
                    prev_char = text[i - 1] if i > 0 else ''
                    next_char = text[i + 1] if i + 1 < len(text) else ''

                    # 前后字符至少有一个是中文，才判定为半角标点误用
                    if ('\u4e00' <= prev_char <= '\u9fff' or
                            '\u4e00' <= next_char <= '\u9fff'):
                        # 排除引用格式中的半角冒号（如 GB/T 1.1:2020）
                        if char == ':' and re.search(r'[A-Z]/[A-Z]\s*\d*$', text[:i]):
                            continue
                        self._add_issue(
                            "W001",
                            Severity.WARNING,
                            f"段落 {para['index']}",
                            f"中文正文使用半角标点\"{char}\"",
                            f"改为全角标点\"{half_to_full[char]}\"",
                            text[:100]
                        )
                        break  # 每段只报一次

    def _check_toc_numbering(self) -> None:
        """F010: 检查目次页码是否使用阿拉伯数字（应为罗马数字）"""
        toc_paras = []
        in_toc = False
        for para in self.paragraphs:
            text = para["text"]
            if "目次" in text and len(text) <= 10:
                in_toc = True
                continue
            if in_toc and para["is_heading"]:
                break
            if in_toc and text:
                toc_paras.append(para)

        # 匹配目次条目末尾的阿拉伯页码
        page_pattern = re.compile(r'[.·…\s]+\s*(\d+)\s*$')

        for para in toc_paras:
            text = para["text"]
            match = page_pattern.search(text)
            if match:
                self._add_issue(
                    "F010",
                    Severity.ERROR,
                    f"段落 {para['index']}",
                    f"目次页码使用阿拉伯数字\"{match.group(1)}\"",
                    "目次页码应使用罗马数字（I、II、III...）",
                    text[:100]
                )

    # ===== 中等难度规则 =====

    def _check_figure_table_footnotes(self) -> None:
        """F012: 检查图/表脚注编号格式（应为字母a/b/c，非数字）"""
        # 匹配"图N"或"表N"附近的脚注使用数字编号
        fig_footnote_pattern = re.compile(r'图\s*\d+.*?注\s*(\d+)')
        tab_footnote_pattern = re.compile(r'表\s*\d+.*?注\s*(\d+)')

        for para in self.paragraphs:
            text = para["text"]
            for pattern in [fig_footnote_pattern, tab_footnote_pattern]:
                match = pattern.search(text)
                if match:
                    self._add_issue(
                        "F012",
                        Severity.ERROR,
                        f"段落 {para['index']}",
                        f"图/表脚注使用数字编号\"{match.group(1)}\"",
                        "图/表脚注应使用字母编号：a、b、c...",
                        text[:100]
                    )

    def _check_dashed_paragraphs(self) -> None:
        """S005: 检查款/项无归属（破折号引导的款直接出现在章下）"""
        dash_pattern = re.compile(r'^[—–-]+\s*')

        for para in self.paragraphs:
            text = para["text"]
            if not dash_pattern.match(text):
                continue

            # 找到最近的标题
            prev_heading = None
            for h in reversed(self.headings):
                if h["index"] <= para["index"]:
                    prev_heading = h
                    break

            if prev_heading:
                heading_text = prev_heading["text"]
                # 如果上一个标题是章标题（数字+空格，无点号）
                first_token = heading_text.split()[0] if heading_text.split() else ""
                if re.match(r'^\d+$', first_token):
                    self._add_issue(
                        "S005",
                        Severity.ERROR,
                        f"段落 {para['index']}",
                        "款直接出现在章下，无归属条",
                        "将款移至某一条下",
                        text[:100]
                    )

    def _check_empty_sections(self) -> None:
        """S007: 检查无内容要素是否已声明"""
        sections_to_check = [
            ("规范性引用文件", "本文件没有规范性引用文件"),
            ("术语和定义", "本文件中没有需要界定的术语和定义"),
        ]

        for section_name, declaration in sections_to_check:
            section_paras = [p for p in self.paragraphs
                             if p.get("section") == section_name]

            if not section_paras:
                self._add_issue(
                    "S007",
                    Severity.ERROR,
                    "整体",
                    f"缺少\"{section_name}\"要素（即使无内容也应列出标题并声明）",
                    f"添加\"{section_name}\"标题，并写明\"{declaration}。\"",
                )
                continue

            content_paras = [p for p in section_paras if not p["is_heading"]]
            if not content_paras:
                self._add_issue(
                    "S007",
                    Severity.ERROR,
                    section_name,
                    f"\"{section_name}\"部分只有标题无内容",
                    f"写明\"{declaration}。\"或列出相关内容",
                )
                continue

            all_text = " ".join(p["text"] for p in content_paras)
            if declaration in all_text:
                continue

            # 检查是否有实际内容
            if section_name == "规范性引用文件":
                has_content = bool(
                    re.search(r'GB[/／]T\s*\d+', all_text) or
                    re.search(r'ISO\s*\d+', all_text) or
                    re.search(r'IEC\s*\d+', all_text)
                )
            else:
                has_content = any(
                    "是指" in p["text"] or "定义" in p["text"]
                    for p in content_paras if len(p["text"]) > 5
                )

            if not has_content:
                self._add_issue(
                    "S007",
                    Severity.ERROR,
                    section_name,
                    f"\"{section_name}\"部分无内容且未声明",
                    f"写明\"{declaration}。\"或列出相关内容",
                )

    def _check_comply_conform(self) -> None:
        """T004: 检查"遵守"和"符合"混用"""
        for para in self.paragraphs:
            text = para["text"]

            # "遵守...要求" → 应改为"符合...要求"
            if re.search(r'遵守.*?要求', text):
                self._add_issue(
                    "T004",
                    Severity.WARNING,
                    f"段落 {para['index']}",
                    "\"遵守\"与\"要求\"搭配不当",
                    "改为\"符合...要求\"（产品/对象符合要求）",
                    text[:100]
                )

            # "符合...规则/原则" → 应改为"遵守...规则/原则"
            if re.search(r'符合.*?规则', text) or re.search(r'符合.*?原则', text):
                self._add_issue(
                    "T004",
                    Severity.WARNING,
                    f"段落 {para['index']}",
                    "\"符合\"与\"规则/原则\"搭配不当",
                    "改为\"遵守...规则/原则\"（人/组织遵守规则）",
                    text[:100]
                )

    def _check_footnote_requirements(self) -> None:
        """T007: 检查条文脚注是否含要求条款"""
        if not self.doc:
            return

        try:
            from lxml import etree
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

            footnotes_part = None
            for rel in self.doc.part.rels.values():
                if 'footnotes' in rel.reltype:
                    footnotes_part = rel.target_part
                    break

            if footnotes_part:
                root = etree.fromstring(footnotes_part.blob)
                for fn in root.findall('.//w:footnote', ns):
                    fn_id = fn.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                    if fn_id in ('-1', '0'):
                        continue
                    texts = fn.findall('.//w:t', ns)
                    fn_text = ''.join(t.text or '' for t in texts)
                    if fn_text and re.search(r'应[\u4e00-\u9fff]', fn_text):
                        self._add_issue(
                            "T007",
                            Severity.ERROR,
                            f"脚注 {fn_id}",
                            "条文脚注含要求条款（出现\"应\"）",
                            "条文脚注只能含陈述性内容",
                            fn_text[:100]
                        )
        except Exception:
            pass

    def _check_term_definition_requirements(self) -> None:
        """T009: 检查术语定义是否含要求型条款"""
        term_paras = [p for p in self.paragraphs
                      if p.get("section") == "术语和定义" and not p["is_heading"]]

        for para in term_paras:
            text = para["text"]
            if ("是指" in text or "指" in text) and re.search(r'应[\u4e00-\u9fff]', text):
                self._add_issue(
                    "T009",
                    Severity.ERROR,
                    f"段落 {para['index']}",
                    "术语定义含要求型条款（出现\"应\"）",
                    "术语定义不应含要求型条款，改为陈述性表述",
                    text[:100]
                )

    def _check_term_bold(self) -> None:
        """W003: 检查术语首次出现是否加粗"""
        if not self.doc:
            return

        term_section_started = False

        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            # 检测进入术语和定义部分
            if self._get_section_name(text) == "术语和定义":
                term_section_started = True
                continue

            if not term_section_started:
                continue

            # 检测离开术语和定义部分
            style_name = para.style.name if para.style else ""
            if self._is_heading(text, style_name):
                if self._get_section_name(text) != "术语和定义":
                    break

            # 检查是否是术语定义段落
            if "是指" not in text and "指" not in text:
                continue

            term_end = text.find("是指") if "是指" in text else text.find("指")
            if term_end <= 0:
                continue

            term_text = text[:term_end].strip()
            if not term_text or len(term_text) > 30:
                continue

            # 检查术语部分的 runs 是否加粗
            bold_found = False
            char_count = 0
            for run in para.runs:
                run_text = run.text or ""
                if char_count >= term_end:
                    break
                if run.bold:
                    bold_found = True
                char_count += len(run_text)

            if not bold_found:
                self._add_issue(
                    "W003",
                    Severity.SUGGESTION,
                    f"段落 {i}",
                    f"术语\"{term_text}\"首次出现未加粗",
                    f"加粗术语\"{term_text}\"",
                    text[:100]
                )


def main():
    parser = argparse.ArgumentParser(
        description="标准草稿自动化检查工具（引用标准符合性检查 v2.6）"
    )
    parser.add_argument("input", nargs="?", help="输入文件路径（.docx 或 .pdf）")
    parser.add_argument("--output", "-o", help="输出 JSON 结果文件路径（可选）")
    parser.add_argument("--pretty", "-p", action="store_true", help="格式化输出 JSON")
    parser.add_argument("--analyze-styles", "-a", action="store_true", help="分析文档样式（仅 .docx）")
    parser.add_argument(
        "--standard", "-s",
        help="指定标准类型（默认自动检测）。可选："
             "gb-national(国家标准), gb-industry(行业标准), "
             "gb-local(地方标准), gb-enterprise(企业标准), gb-group(团体标准)"
    )
    parser.add_argument(
        "--list-standards", action="store_true",
        help="列出所有支持的标准类型"
    )
    parser.add_argument(
        "--ref", action="append", default=[],
        help="引用标准文档路径（用于符合性检查，可多次指定）。如 --ref GB-T1.1-2020.docx"
    )
    parser.add_argument(
        "--ref-dir",
        help="引用标准文档目录路径（自动加载目录下所有 .docx 和 .pdf 文件）"
    )

    args = parser.parse_args()

    # --list-standards：列出标准类型
    if args.list_standards:
        print("支持的标准类型：\n")
        for p in list_profiles():
            print(f"  {p['id']:16s}  {p['name']}  ({p['number_example']})")
            print(f"  {'':16s}  {p['description']}\n")
        return

    if not args.input:
        parser.print_help()
        sys.exit(1)

    # 检查输入文件
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"错误：文件不存在：{input_path}")
        sys.exit(1)

    file_ext = input_path.suffix.lower()
    if file_ext not in (".docx", ".pdf"):
        print(f"错误：不支持的文件格式：{file_ext}（仅支持 .docx 和 .pdf）")
        sys.exit(1)

    # 验证 --standard 参数
    if args.standard:
        try:
            get_profile(args.standard)
        except ValueError as e:
            print(f"错误：{e}")
            sys.exit(1)

    # 执行检查
    checker = StandardChecker(standard=args.standard)

    # 收集引用标准文件
    ref_files = list(args.ref) if args.ref else []
    if args.ref_dir:
        ref_dir = Path(args.ref_dir)
        if not ref_dir.exists():
            print(f"错误：引用标准目录不存在：{ref_dir}")
            sys.exit(1)
        for ext in ("*.docx", "*.pdf"):
            for f in ref_dir.glob(ext):
                ref_files.append(str(f))
        if ref_files:
            print(f"从 {ref_dir} 加载了 {len(ref_files)} 个引用标准文件")

    issues = checker.check(str(input_path), analyze_styles=args.analyze_styles,
                           ref_files=ref_files if ref_files else None)
    
    # 输出结果
    result = {
        "file": str(input_path),
        "standard_type": checker.profile.id if checker.profile else "gb-national",
        "standard_name": checker.profile.name if checker.profile else "国家标准",
        "drafting_standard": checker.profile.drafting_standard if checker.profile else "GB/T 1.1-2020",
        "total_issues": len(issues),
        "summary": {
            "ERROR": sum(1 for i in issues if i.severity == "ERROR"),
            "WARNING": sum(1 for i in issues if i.severity == "WARNING"),
            "SUGGESTION": sum(1 for i in issues if i.severity == "SUGGESTION")
        },
        "issues": [asdict(i) for i in issues]
    }
    
    # 添加引用标准检查摘要
    if hasattr(checker, 'ref_checker'):
        result["reference_check"] = checker.ref_checker.get_summary()
    
    # 输出到文件或控制台
    if args.output:
        output_path = Path(args.output)
        indent = 2 if args.pretty else None
        output_path.write_text(json.dumps(result, indent=indent, ensure_ascii=False), encoding="utf-8")
        print(f"检查完成，结果已保存到：{output_path}")
    else:
        indent = 2 if args.pretty else None
        print(json.dumps(result, indent=indent, ensure_ascii=False))
    
    # 返回退出码
    if result["summary"]["ERROR"] > 0:
        sys.exit(1)
    else:
        sys.exit(0)


if __name__ == "__main__":
    main()